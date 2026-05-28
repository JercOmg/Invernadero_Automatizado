"""
generar_documentacion_en.py
Generates the complete documentation of the Automated Greenhouse Management System in English (.docx).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────
# Style helpers
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '2E7D32')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_header_row(table, headers, bg='2E7D32', fg='FFFFFF'):
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(fg)
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, bg)
        set_cell_border(cell)

def add_data_row(table, row_index, values, alternate=False):
    row = table.rows[row_index]
    bg = 'E8F5E9' if alternate else 'FFFFFF'
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(str(val))
        run.font.size = Pt(8.5)
        set_cell_bg(cell, bg)
        set_cell_border(cell)

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0] if heading.runs else heading.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
        run.font.size = Pt(13)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x38, 0x8E, 0x3C)
        run.font.size = Pt(11)
    return heading

def add_paragraph(doc, text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x01, 0x57, 0x9B)

def page_break(doc):
    doc.add_page_break()


# ─────────────────────────────────────────────
# DOCUMENT
# ─────────────────────────────────────────────

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# Base font style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ═══════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AUTOMATED GREENHOUSE\nMANAGEMENT SYSTEM')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('Technical Project Documentation')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f'Version: 1.0.0\n').font.size = Pt(11)
info.add_run(f'Date: {datetime.date.today().strftime("%B %d, %Y")}\n').font.size = Pt(11)
info.add_run('Course: Software Factory').font.size = Pt(11)

page_break(doc)

# ═══════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════
add_heading(doc, 'Table of Contents', level=1)

toc_items = [
    ('1.', 'Introduction'),
    ('2.', 'Problem Statement'),
    ('3.', 'Objectives'),
    ('4.', 'State of the Art – Background and Related Work'),
    ('5.', 'Requirements'),
    ('6.', 'Use Cases and User Stories'),
    ('7.', 'Data Dictionary and Entity-Relationship Model'),
    ('8.', 'Class Diagram'),
    ('9.', 'User Interface Design'),
    ('10.', 'Web Services Catalog and API Documentation'),
    ('11.', 'Testing (Unit, Functional and Integration)'),
    ('12.', 'Proposed Architecture Diagram'),
    ('13.', 'Results and Discussion'),
]

for num, item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'{num}  {item}')
    run.font.size = Pt(11)

page_break(doc)

# ═══════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════
add_heading(doc, '1. Introduction', level=1)
add_paragraph(doc,
    'This document describes the design, development, and implementation of the Automated Greenhouse '
    'Management System, a web application developed as an academic project for the Software Factory '
    'course. The system aims to centralize and automate the critical operational processes of a '
    'greenhouse: environmental monitoring via sensors, crop and harvest management, irrigation control, '
    'supply administration, and alert generation for anomalous conditions.'
)
add_paragraph(doc,
    'The solution is built on a three-tier architecture: a REST backend developed with Spring Boot 3.2 '
    'and Java 17, a Single Page Application (SPA) frontend developed with React 19 and Vite 6, and a '
    'relational database on PostgreSQL. The project also incorporates automation through Python scripts '
    'for generating artifacts (database model, data dictionary, backend structure, and frontend), a '
    'CI/CD pipeline with GitHub Actions, Google OAuth 2.0 authentication, internationalization, and '
    'testing across all system layers.'
)
add_paragraph(doc,
    'One of the distinguishing aspects of this project is its focus on automating the development '
    'process itself: the data model is defined once in a JSON file, and from it the SQL schema, '
    'data dictionary, REST controllers, and React CRUD pages are all generated automatically.'
)

page_break(doc)

# ═══════════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ═══════════════════════════════════════════════
add_heading(doc, '2. Problem Statement', level=1)

add_heading(doc, '2.1 Problem Description', level=2)
add_paragraph(doc,
    'Modern greenhouses face significant challenges in manually managing their operations. '
    'Operators must record temperature, humidity, pH, and luminosity readings by hand, introducing '
    'human error and delays in detecting critical conditions that can compromise crops. Additionally, '
    'supply inventory management, irrigation records, and crop lifecycle tracking are often maintained '
    'in spreadsheets or on paper, making traceability and analysis difficult.'
)

add_heading(doc, '2.2 Identified Root Causes', level=2)
add_bullet(doc, 'Absence of a centralized system for monitoring environmental variables.')
add_bullet(doc, 'Manual processes prone to human error in operational record-keeping.')
add_bullet(doc, 'Lack of automatic alerts when readings exceed optimal thresholds.')
add_bullet(doc, 'Difficulty in achieving end-to-end traceability of complete crop cycles.')
add_bullet(doc, 'No multi-role access control with differentiated permission levels.')
add_bullet(doc, 'Absence of real-time dashboards and visualizations of greenhouse status.')

add_heading(doc, '2.3 Impact of the Problem', level=2)
add_paragraph(doc,
    'The lack of automation leads to crop losses due to delayed responses to adverse environmental '
    'conditions, supply waste due to poor inventory control, and difficulty in identifying patterns '
    'and optimizing the productive processes of the greenhouse.'
)

page_break(doc)

# ═══════════════════════════════════════════════
# 3. OBJECTIVES
# ═══════════════════════════════════════════════
add_heading(doc, '3. Objectives', level=1)

add_heading(doc, '3.1 General Objective', level=2)
add_paragraph(doc,
    'Develop a comprehensive web-based greenhouse management system that automates environmental '
    'variable monitoring, irrigation control, crop management, and supply administration, integrating '
    'a secure, internationalized REST architecture with continuous deployment and test coverage across '
    'all layers.'
)

add_heading(doc, '3.2 Specific Objectives', level=2)
specifics = [
    'Design and automatically generate the relational database model in PostgreSQL from a JSON source-of-truth model.',
    'Implement a REST backend with Spring Boot 3.2 exposing CRUD services for all 12 entities, with JWT-based security and Google OAuth 2.0 authentication.',
    'Develop an SPA frontend with React 19 that consumes the REST services and supports internationalization (Spanish and English).',
    'Configure a CI/CD pipeline with GitHub Actions that runs automated tests (JUnit, Selenium, pytest) and creates tasks in Taiga on failure.',
    'Implement an alert system that detects anomalous sensor readings and notifies the appropriate users.',
    'Document the REST API using SpringDoc OpenAPI (Swagger UI) with complete annotations on all endpoints.',
    'Achieve test coverage exceeding 80% on the backend and cover critical frontend flows with E2E tests (Selenium).',
    'Automate the generation of system components (backend and frontend) through Python scripts.',
]
for obj in specifics:
    add_bullet(doc, obj)

page_break(doc)

# ═══════════════════════════════════════════════
# 4. STATE OF THE ART
# ═══════════════════════════════════════════════
add_heading(doc, '4. State of the Art – Background and Related Work', level=1)

add_heading(doc, '4.1 Smart Greenhouses and IoT', level=2)
add_paragraph(doc,
    'The concept of smart greenhouses has evolved significantly over the past decade, driven by the '
    'Internet of Things (IoT). Systems such as Greenhouse Manager (Priva), iGrow, and Argus Controls '
    'integrate temperature, relative humidity, CO₂, and pH sensors with automatic control systems for '
    'ventilation, irrigation, and fertilization. However, these commercial solutions carry high '
    'licensing costs and limited flexibility for small and medium-sized production units.'
)

add_heading(doc, '4.2 Related Academic Work', level=2)
related = [
    ('IoT Monitoring with Arduino and MQTT (2022)',
     'Low-cost systems based on Arduino + Raspberry Pi using the MQTT protocol to transmit sensor '
     'readings to platforms like ThingsBoard. These validated the technical feasibility of remote '
     'real-time monitoring with hardware costs under USD 100.'),
    ('Automated Irrigation Systems with Machine Learning (2023)',
     'Research combining soil moisture sensors with ML models (Random Forest, Neural Networks) for '
     'optimal irrigation timing prediction, achieving reductions in water consumption of up to 30% '
     'compared to timer-based irrigation systems.'),
    ('Web Platforms for Agricultural Management (2021–2024)',
     'Academic projects implementing REST + React architectures for crop management, highlighting '
     'Spring Boot as the backend of choice due to its mature ecosystem, security support (Spring '
     'Security), and ease of documentation (SpringDoc/Swagger).'),
    ('Code Generation from Models (2024)',
     'Software factory projects implementing the code generation pattern from JSON models, reducing '
     'repetitive development effort and ensuring consistency between the data model and implementation.'),
]
for title, desc in related:
    add_paragraph(doc, f'● {title}', bold=True)
    add_paragraph(doc, f'  {desc}')

add_heading(doc, '4.3 Technology Stack and Justification', level=2)
tech_rows = [
    ('Spring Boot 3.2 / Java 17', 'Mature REST API framework with integrated security, JPA, and documentation ecosystems.'),
    ('React 19 + Vite 6', 'High-performance SPA with a robust ecosystem for i18n and OAuth.'),
    ('PostgreSQL', 'Robust relational database with advanced type support and ACID transactions.'),
    ('Python 3.10+', 'Automation scripts for artifact generation and pytest-based testing.'),
    ('GitHub Actions', 'Native CI/CD for GitHub, free for public and private repositories.'),
    ('Google OAuth 2.0', 'Delegated authentication to a trusted identity provider — no stored passwords.'),
    ('JWT (JJWT 0.12.5)', 'Stateless tokens for authorization, scalable and compatible with distributed architectures.'),
    ('SpringDoc OpenAPI 2.5', 'Automatic generation of Swagger UI documentation from Java annotations.'),
    ('Selenium + WebDriver', 'E2E frontend testing in a real browser — industry standard.'),
    ('Taiga', 'Agile project management platform integrated with CI/CD for traceability.'),
]
table_tech = doc.add_table(rows=len(tech_rows)+1, cols=2)
table_tech.style = 'Table Grid'
table_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_tech, ['Technology', 'Justification'])
for i, row_data in enumerate(tech_rows):
    add_data_row(table_tech, i+1, row_data, alternate=(i % 2 == 0))

page_break(doc)

# ═══════════════════════════════════════════════
# 5. REQUIREMENTS
# ═══════════════════════════════════════════════
add_heading(doc, '5. Requirements', level=1)

add_heading(doc, '5.1 Functional Requirements', level=2)

rf_data = [
    ('FR-01', 'User Management', 'The system must allow registration, local and Google OAuth authentication, editing, and deactivation of users with differentiated roles (ADMINISTRATOR, TECHNICIAN, OPERATOR, VIEWER).', 'High'),
    ('FR-02', 'Greenhouse Management', 'The system must allow creating, listing, updating, and deleting greenhouses with attributes: name, location, area, and structure type.', 'High'),
    ('FR-03', 'Zone Management', 'The system must allow defining internal zones within each greenhouse.', 'High'),
    ('FR-04', 'Crop Catalog', 'The system must allow registering and maintaining a crop catalog with optimal parameters (temperature, humidity, cycle length).', 'High'),
    ('FR-05', 'Planting Records', 'The system must record planting events associating zone, crop, user, and lifecycle status tracking.', 'High'),
    ('FR-06', 'Sensor Management', 'The system must allow registering and monitoring sensors per zone. Types: TEMPERATURE, HUMIDITY, CO2, LUMINOSITY, PH, SOIL_HUMIDITY.', 'High'),
    ('FR-07', 'Sensor Readings', 'The system must store the time series of readings per sensor and flag those exceeding defined thresholds.', 'High'),
    ('FR-08', 'Irrigation Control', 'The system must record irrigation events (automatic or manual) with duration and volume.', 'Medium'),
    ('FR-09', 'Alert System', 'The system must generate and notify alerts classified by level (INFORMATIONAL, WARNING, CRITICAL) for anomalous conditions.', 'High'),
    ('FR-10', 'Supply Management', 'The system must manage supply inventory with minimum stock control and replenishment alerts.', 'Medium'),
    ('FR-11', 'Supply Application', 'The system must record supply applications to plantings or zones with full traceability.', 'Medium'),
    ('FR-12', 'Harvest Records', 'The system must record harvested production per planting event with quality classification.', 'High'),
    ('FR-13', 'Dashboard', 'The system must display a control panel with key greenhouse status metrics.', 'High'),
    ('FR-14', 'Internationalization', 'The system must support the interface in Spanish and English with real-time switching.', 'Medium'),
    ('FR-15', 'Taiga Integration', 'The CI/CD system must automatically create tasks in Taiga when tests fail.', 'Medium'),
    ('FR-16', 'API Documentation', 'All REST endpoints must be documented and accessible via Swagger UI.', 'High'),
]

table_rf = doc.add_table(rows=len(rf_data)+1, cols=4)
table_rf.style = 'Table Grid'
table_rf.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_rf, ['ID', 'Name', 'Description', 'Priority'])
for i, row_data in enumerate(rf_data):
    add_data_row(table_rf, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '5.2 Non-Functional Requirements', level=2)

rnf_data = [
    ('NFR-01', 'Security', 'JWT authentication with short-lived tokens and Google ID Token validation. Passwords stored with BCrypt.'),
    ('NFR-02', 'Performance', 'List queries must respond in under 500ms under normal load. Pagination available on all list endpoints.'),
    ('NFR-03', 'Availability', 'The system must guarantee a minimum uptime of 99% in the production environment (Render.com).'),
    ('NFR-04', 'Scalability', 'Stateless architecture (JWT) allows horizontal scaling without additional configuration.'),
    ('NFR-05', 'Maintainability', 'Code documented with Javadoc/docstrings. Test coverage > 80%. Layered architecture (controller → service → repository).'),
    ('NFR-06', 'Portability', 'The backend runs on any environment with Java 17+. The frontend is a static bundle deployable to any CDN.'),
    ('NFR-07', 'Usability', 'Responsive interface with multilingual support and visual feedback on all CRUD operations.'),
    ('NFR-08', 'Traceability', 'All critical events (planting, harvest, alerts, applications) are recorded with timestamp and responsible user.'),
]
table_rnf = doc.add_table(rows=len(rnf_data)+1, cols=3)
table_rnf.style = 'Table Grid'
table_rnf.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_rnf, ['ID', 'Name', 'Description'])
for i, row_data in enumerate(rnf_data):
    add_data_row(table_rnf, i+1, row_data, alternate=(i % 2 == 0))

page_break(doc)

# ═══════════════════════════════════════════════
# 6. USE CASES AND USER STORIES
# ═══════════════════════════════════════════════
add_heading(doc, '6. Use Cases and User Stories', level=1)

add_heading(doc, '6.1 System Actors', level=2)
actors = [
    ('Administrator', 'Full system access. Manages users, greenhouses, configurations, and views all reports.'),
    ('Technician', 'Manages sensors, zones, irrigation, and generates alerts. Can register plantings and harvests.'),
    ('Operator', 'Records manual readings, applies supplies, and executes irrigation. Limited visibility to their assigned zone.'),
    ('Viewer', 'Read-only access to dashboards and reports.'),
    ('System (Automatic)', 'Generates threshold alerts and creates tasks in Taiga on CI/CD failures.'),
]
for actor, desc in actors:
    add_paragraph(doc, f'● {actor}:', bold=True)
    add_paragraph(doc, f'  {desc}')

add_heading(doc, '6.2 Main Use Cases', level=2)

cu_data = [
    ('UC-01', 'Authenticate User', 'Operator, Technician, Admin', 'User with valid credentials or Google account', 'User authenticated with JWT token'),
    ('UC-02', 'Register Greenhouse', 'Administrator', 'Greenhouse data: name, location, area, type', 'Greenhouse saved in database'),
    ('UC-03', 'Record Sensor Reading', 'Technician, System', 'Active sensor and measured value', 'Reading stored; alert generated if applicable'),
    ('UC-04', 'Manage Planting', 'Technician, Operator', 'Zone, crop, and planting data', 'Planting cycle started with traceability'),
    ('UC-05', 'Execute Irrigation', 'Operator, System', 'Zone and irrigation parameters', 'Irrigation event recorded'),
    ('UC-06', 'Manage Alerts', 'Technician, Administrator', 'List of active alerts', 'Alerts resolved with resolution timestamp'),
    ('UC-07', 'Record Harvest', 'Technician, Operator', 'Completed planting and production data', 'Harvest recorded with quantity and quality'),
    ('UC-08', 'View Dashboard', 'All roles', 'Active session', 'View of metrics and general greenhouse status'),
    ('UC-09', 'Manage Supplies', 'Technician, Administrator', 'Supply data and stock levels', 'Inventory updated with minimum stock alerts'),
    ('UC-10', 'Change Language', 'All roles', 'Language selection (ES/EN)', 'Interface translated in real time'),
]

table_cu = doc.add_table(rows=len(cu_data)+1, cols=5)
table_cu.style = 'Table Grid'
table_cu.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_cu, ['ID', 'Use Case', 'Actor(s)', 'Precondition', 'Postcondition'])
for i, row_data in enumerate(cu_data):
    add_data_row(table_cu, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '6.3 User Stories', level=2)

hu_data = [
    ('US-01', 'Authentication', 'As a system user, I want to sign in with my Google account so that I do not have to remember another password.',
     '● Must display a "Sign in with Google" button\n● Upon authentication, redirects to Dashboard\n● JWT token expires after 24 hours\n● If the email does not exist, the account is created automatically'),
    ('US-02', 'Sensor Monitoring', 'As a technician, I want to view real-time readings from all sensors so that I can detect out-of-range conditions.',
     '● Lists all readings with sensor, value, and timestamp\n● Highlights in red the readings that triggered an alert\n● Allows filtering by zone and sensor type'),
    ('US-03', 'Critical Alerts', 'As an administrator, I want to receive alerts classified by severity level so that I can prioritize my response.',
     '● Alerts are classified as INFORMATIONAL, WARNING, and CRITICAL\n● An alert can be marked as resolved\n● Resolution date and time are recorded'),
    ('US-04', 'Planting Cycle', 'As an operator, I want to register a new planting indicating the zone, crop, and number of plants to maintain lifecycle traceability.',
     '● Required fields: zone, crop, planting date, number of plants\n● Initial status is IN_GROWTH\n● Estimated harvest date is calculated and displayed'),
    ('US-05', 'Supply Inventory', 'As a technician, I want to see an alert when a supply stock falls below the minimum so that I can request restocking in time.',
     '● The supply table highlights in red items below minimum stock\n● New inventory entries can be recorded\n● Application history automatically reduces stock'),
    ('US-06', 'Internationalization', 'As a user, I want to switch the interface language between Spanish and English at any time.',
     '● Language change is immediate without page reload\n● Preference is saved in localStorage\n● All texts, error messages, and labels are translated'),
    ('US-07', 'Harvest Record', 'As a technician, I want to record harvested production with quantity in kg and quality classification for full traceability.',
     '● Associated with an existing planting in IN_GROWTH or HARVESTED status\n● Quality options: PREMIUM, STANDARD, SECOND, DISCARD\n● Recording a harvest changes planting status to HARVESTED'),
    ('US-08', 'Executive Dashboard', 'As an administrator, I want to see a greenhouse status summary on the dashboard so that I can make informed decisions quickly.',
     '● Shows total greenhouses, active zones, pending alerts, and ongoing plantings\n● Charts of latest sensor readings per zone\n● Quick access to most-used sections'),
]

for hu in hu_data:
    add_heading(doc, f'{hu[0]}: {hu[1]}', level=3)
    add_paragraph(doc, hu[2], italic=True)
    add_paragraph(doc, 'Acceptance Criteria:', bold=True)
    for line in hu[3].split('\n'):
        if line.strip():
            add_bullet(doc, line.strip().lstrip('●').strip())
    doc.add_paragraph()

page_break(doc)

# ═══════════════════════════════════════════════
# 7. DATA DICTIONARY AND ER MODEL
# ═══════════════════════════════════════════════
add_heading(doc, '7. Data Dictionary and Entity-Relationship Model', level=1)

add_heading(doc, '7.1 Model Summary', level=2)
add_paragraph(doc,
    'The data model of the Automated Greenhouse Management System consists of 12 entities '
    'that capture all operational processes of the greenhouse. The model was automatically '
    'generated from the base_datos_invernadero.json file using the Python script '
    'generar_base_de_datos.py.'
)

entities_summary = [
    ('USER', 9, 'System users with roles and access permissions'),
    ('GREENHOUSE', 8, 'General information about each physical greenhouse installation'),
    ('ZONE', 5, 'Internal sections or areas within a greenhouse'),
    ('CROP', 10, 'Master catalog of species and varieties cultivated'),
    ('PLANTING', 9, 'Record of each planting event in a zone'),
    ('SENSOR', 7, 'Measurement devices installed in zones'),
    ('SENSOR_READING', 5, 'Time series of measurements recorded by each sensor'),
    ('IRRIGATION', 8, 'Automatic or manual irrigation events per zone'),
    ('ALERT', 9, 'Notifications for anomalous conditions detected'),
    ('SUPPLY', 7, 'Supply catalog with inventory control'),
    ('SUPPLY_APPLICATION', 9, 'History of supply applications to plantings or zones'),
    ('HARVEST', 7, 'Record of production harvested per planting event'),
]

table_ent = doc.add_table(rows=len(entities_summary)+1, cols=3)
table_ent.style = 'Table Grid'
table_ent.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_ent, ['Entity', 'Fields', 'Description'])
for i, row_data in enumerate(entities_summary):
    add_data_row(table_ent, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '7.2 Data Dictionary by Entity', level=2)

entities_dict = {
    'USER': [
        ('id_usuario', 'INT', 'NO', 'YES', 'NO', 'NO', '—', 'Unique user identifier'),
        ('nombre', 'VARCHAR(100)', 'NO', 'NO', 'NO', 'NO', '—', 'First name of the user'),
        ('apellido', 'VARCHAR(100)', 'NO', 'NO', 'NO', 'NO', '—', 'Last name of the user'),
        ('email', 'VARCHAR(150)', 'NO', 'NO', 'NO', 'YES', '—', 'Unique email used for OAuth authentication'),
        ('password_hash', 'VARCHAR(255)', 'YES', 'NO', 'NO', 'NO', '—', 'Password hash (null if using Google OAuth only)'),
        ('rol', 'VARCHAR(30)', 'NO', 'NO', 'NO', 'NO', 'OPERATOR', 'Role: ADMINISTRATOR|TECHNICIAN|OPERATOR|VIEWER'),
        ('fecha_registro', 'DATE', 'NO', 'NO', 'NO', 'NO', '—', 'Date the user was registered in the system'),
        ('activo', 'BOOLEAN', 'NO', 'NO', 'NO', 'NO', 'True', 'Indicates whether the account is active'),
        ('proveedor_oauth', 'VARCHAR(30)', 'YES', 'NO', 'NO', 'NO', '—', 'OAuth provider: LOCAL|GOOGLE'),
    ],
    'GREENHOUSE': [
        ('id_invernadero', 'INT', 'NO', 'YES', 'NO', 'NO', '—', 'Unique greenhouse identifier'),
        ('nombre', 'VARCHAR(100)', 'NO', 'NO', 'NO', 'NO', '—', 'Name or code of the greenhouse'),
        ('ubicacion', 'VARCHAR(200)', 'NO', 'NO', 'NO', 'NO', '—', 'Physical location description'),
        ('area_m2', 'DECIMAL(10,2)', 'NO', 'NO', 'NO', 'NO', '—', 'Total area in square meters'),
        ('tipo_estructura', 'VARCHAR(50)', 'YES', 'NO', 'NO', 'NO', '—', 'Structure material: GLASS|POLYCARBONATE|MESH|PLASTIC'),
        ('responsable_id', 'INT', 'NO', 'NO', 'YES', 'NO', '—', 'FK → user.id_usuario'),
        ('fecha_creacion', 'DATE', 'NO', 'NO', 'NO', 'NO', '—', 'Date the greenhouse was registered'),
        ('estado', 'VARCHAR(20)', 'NO', 'NO', 'NO', 'NO', 'ACTIVE', 'Status: ACTIVE|INACTIVE|MAINTENANCE'),
    ],
    'SENSOR': [
        ('id_sensor', 'INT', 'NO', 'YES', 'NO', 'NO', '—', 'Unique sensor identifier'),
        ('id_zona', 'INT', 'NO', 'NO', 'YES', 'NO', '—', 'FK → zone.id_zona'),
        ('tipo_sensor', 'VARCHAR(50)', 'NO', 'NO', 'NO', 'NO', '—', 'TEMPERATURE|HUMIDITY|CO2|LUMINOSITY|PH|SOIL_HUMIDITY'),
        ('modelo', 'VARCHAR(100)', 'YES', 'NO', 'NO', 'NO', '—', 'Commercial model or reference'),
        ('unidad_medida', 'VARCHAR(20)', 'NO', 'NO', 'NO', 'NO', '—', 'Unit of measurement (°C, %, ppm, lux)'),
        ('fecha_instalacion', 'DATE', 'YES', 'NO', 'NO', 'NO', '—', 'Installation date'),
        ('estado', 'VARCHAR(20)', 'NO', 'NO', 'NO', 'NO', 'ACTIVE', 'Status: ACTIVE|INACTIVE|FAULT'),
    ],
    'ALERT': [
        ('id_alerta', 'INT', 'NO', 'YES', 'NO', 'NO', '—', 'Unique alert identifier'),
        ('id_sensor', 'INT', 'YES', 'NO', 'YES', 'NO', '—', 'FK → sensor.id_sensor (null if manual alert)'),
        ('id_zona', 'INT', 'YES', 'NO', 'YES', 'NO', '—', 'FK → zone.id_zona'),
        ('tipo_alerta', 'VARCHAR(50)', 'NO', 'NO', 'NO', 'NO', '—', 'TEMP_HIGH|HUMIDITY_LOW|CO2_HIGH|PH_OUT_OF_RANGE|...'),
        ('descripcion', 'TEXT', 'NO', 'NO', 'NO', 'NO', '—', 'Detailed message describing the alert condition'),
        ('fecha_hora', 'DATETIME', 'NO', 'NO', 'NO', 'NO', '—', 'Alert generation timestamp'),
        ('nivel', 'VARCHAR(20)', 'NO', 'NO', 'NO', 'NO', '—', 'Level: INFORMATIONAL|WARNING|CRITICAL'),
        ('resuelta', 'BOOLEAN', 'NO', 'NO', 'NO', 'NO', 'False', 'Indicates whether the alert has been resolved'),
        ('fecha_resolucion', 'DATETIME', 'YES', 'NO', 'NO', 'NO', '—', 'Resolution timestamp'),
    ],
}

for idx, (entity, fields) in enumerate(entities_dict.items()):
    add_heading(doc, f'7.2.{idx+1} Entity: {entity}', level=3)
    t = doc.add_table(rows=len(fields)+1, cols=8)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(t, ['Field', 'Type', 'Null', 'PK', 'FK', 'Unique', 'Default', 'Description'])
    for i, row_data in enumerate(fields):
        add_data_row(t, i+1, row_data, alternate=(i % 2 == 0))
    doc.add_paragraph()

add_paragraph(doc,
    'Note: The entities ZONE, CROP, PLANTING, SENSOR_READING, IRRIGATION, SUPPLY, '
    'SUPPLY_APPLICATION, and HARVEST follow the same documentation pattern. '
    'The complete dictionary is available in Base de Datos/diccionario_datos.txt '
    'and Base de Datos/diccionario_datos.pdf.',
    italic=True
)

doc.add_paragraph()
add_heading(doc, '7.3 Entity-Relationship Model – Relationship Map', level=2)
add_paragraph(doc, 'The following foreign key relationships define the relational model of the system:')

relationships = [
    'GREENHOUSE.responsable_id      →  USER.id_usuario',
    'ZONE.id_invernadero            →  GREENHOUSE.id_invernadero',
    'PLANTING.id_zona               →  ZONE.id_zona',
    'PLANTING.id_cultivo            →  CROP.id_cultivo',
    'PLANTING.id_usuario            →  USER.id_usuario',
    'SENSOR.id_zona                 →  ZONE.id_zona',
    'SENSOR_READING.id_sensor       →  SENSOR.id_sensor',
    'IRRIGATION.id_zona             →  ZONE.id_zona',
    'IRRIGATION.id_usuario          →  USER.id_usuario',
    'ALERT.id_sensor                →  SENSOR.id_sensor',
    'ALERT.id_zona                  →  ZONE.id_zona',
    'SUPPLY_APPLICATION.id_insumo   →  SUPPLY.id_insumo',
    'SUPPLY_APPLICATION.id_siembra  →  PLANTING.id_siembra',
    'SUPPLY_APPLICATION.id_zona     →  ZONE.id_zona',
    'SUPPLY_APPLICATION.id_usuario  →  USER.id_usuario',
    'HARVEST.id_siembra             →  PLANTING.id_siembra',
    'HARVEST.id_usuario             →  USER.id_usuario',
]
for rel in relationships:
    add_code_block(doc, rel)

page_break(doc)

# ═══════════════════════════════════════════════
# 8. CLASS DIAGRAM
# ═══════════════════════════════════════════════
add_heading(doc, '8. Class Diagram', level=1)

add_paragraph(doc,
    'The backend implements a layered architecture with clear separation of concerns. '
    'Each entity in the data model has a corresponding Java class organized in the following '
    'packages: domain/model (JPA entities), domain/repository (Spring Data interfaces), '
    'application (business services), and interfaces (REST controllers).'
)

add_heading(doc, '8.1 Backend Package Structure', level=2)
structure = [
    'com.invernadero.invernadero_backend',
    '  ├── auth/',
    '  │   ├── application/        → AuthService, JwtService, GoogleTokenService',
    '  │   ├── domain/',
    '  │   │   ├── model/          → Usuario.java',
    '  │   │   ├── repository/     → UsuarioRepository.java',
    '  │   │   └── service/        → UserDetailsServiceImpl.java',
    '  │   └── interfaces/         → AuthController.java',
    '  ├── invernadero/            → (Greenhouse entity)',
    '  ├── zona/                   → (Zone entity)',
    '  ├── cultivo/                → (Crop entity)',
    '  ├── siembra/                → (Planting entity)',
    '  ├── sensor/                 → (Sensor entity)',
    '  ├── lectura_sensor/         → (SensorReading entity)',
    '  ├── riego/                  → (Irrigation entity)',
    '  ├── alerta/                 → (Alert entity)',
    '  ├── insumo/                 → (Supply entity)',
    '  ├── aplicacion_insumo/      → (SupplyApplication entity)',
    '  ├── cosecha/                → (Harvest entity)',
    '  └── shared/',
    '      ├── application/        → Generic response classes',
    '      └── domain/             → Exceptions and utilities',
]
for line in structure:
    add_code_block(doc, line)

doc.add_paragraph()
add_heading(doc, '8.2 Main Classes', level=2)

classes_data = [
    ('Usuario', 'auth/domain/model', '@Entity JPA', 'id, nombre, apellido, email, passwordHash, rol, activo, proveedorOauth'),
    ('Invernadero', 'invernadero/domain/model', '@Entity JPA', 'id, nombre, ubicacion, areaM2, tipoEstructura, responsable(FK), estado'),
    ('Zona', 'zona/domain/model', '@Entity JPA', 'id, invernadero(FK), nombreZona, areaM2, descripcion'),
    ('Sensor', 'sensor/domain/model', '@Entity JPA', 'id, zona(FK), tipoSensor, modelo, unidadMedida, estado'),
    ('LecturaSensor', 'lectura_sensor/domain/model', '@Entity JPA', 'id, sensor(FK), valor, fechaHora, generaAlerta'),
    ('Alerta', 'alerta/domain/model', '@Entity JPA', 'id, sensor(FK), zona(FK), tipoAlerta, nivel, resuelta, fechaResolucion'),
    ('AuthService', 'auth/application', 'Service', 'register(), login(), validateGoogleToken(), generateJwt()'),
    ('JwtService', 'auth/application', 'Service', 'generateToken(), validateToken(), extractEmail()'),
    ('InvernaderoController', 'invernadero/interfaces', '@RestController', 'create(), list(), findById(), update(), delete()'),
    ('SecurityConfig', 'shared/application', '@Configuration', 'Spring Security config, CORS, JWT filters'),
]

table_cl = doc.add_table(rows=len(classes_data)+1, cols=4)
table_cl.style = 'Table Grid'
table_cl.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_cl, ['Class', 'Package', 'Stereotype', 'Main Attributes / Methods'])
for i, row_data in enumerate(classes_data):
    add_data_row(table_cl, i+1, row_data, alternate=(i % 2 == 0))

page_break(doc)

# ═══════════════════════════════════════════════
# 9. UI DESIGN
# ═══════════════════════════════════════════════
add_heading(doc, '9. User Interface Design', level=1)

add_heading(doc, '9.1 UI Technologies', level=2)
add_bullet(doc, 'React 19 + Vite 6: High-performance SPA construction')
add_bullet(doc, 'Tailwind CSS 4: Utility-first styling for responsive design')
add_bullet(doc, 'React Router DOM v7: Client-side routing')
add_bullet(doc, 'react-i18next: Real-time internationalization (ES/EN)')
add_bullet(doc, '@react-oauth/google: Google OAuth button component')
add_bullet(doc, 'Axios: HTTP client for REST API consumption')

add_heading(doc, '9.2 Frontend Route Structure', level=2)

routes_data = [
    ('/login', 'Public', 'Login form with email/password and Google OAuth button'),
    ('/register', 'Public', 'New user registration form'),
    ('/dashboard', 'Protected', 'Main panel with metrics and quick access to modules'),
    ('/invernadero', 'Protected', 'Greenhouse listing and CRUD with modals'),
    ('/zona', 'Protected', 'Zone listing and CRUD'),
    ('/cultivo', 'Protected', 'Crop catalog with optimal parameters'),
    ('/siembra', 'Protected', 'Planting cycle management'),
    ('/sensor', 'Protected', 'Sensor registration and monitoring'),
    ('/lectura_sensor', 'Protected', 'Reading time series with visual indicators'),
    ('/riego', 'Protected', 'Irrigation event history and control'),
    ('/alerta', 'Protected', 'Alert management with level classification'),
    ('/insumo', 'Protected', 'Supply inventory with stock alerts'),
    ('/aplicacion_insumo', 'Protected', 'Supply application history'),
    ('/cosecha', 'Protected', 'Production records per planting'),
]

table_routes = doc.add_table(rows=len(routes_data)+1, cols=3)
table_routes.style = 'Table Grid'
table_routes.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_routes, ['Route', 'Access', 'Description'])
for i, row_data in enumerate(routes_data):
    add_data_row(table_routes, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '9.3 UI Design Patterns', level=2)
patterns = [
    ('Fixed Sidebar Layout', 'Persistent side navigation with icons and module labels, responsive on mobile.'),
    ('List Page + CRUD Modal', 'All entities follow the pattern: paginated table + action buttons + form modal.'),
    ('ProtectedRoute HOC', 'Higher-order component that redirects to /login if no valid JWT token is found.'),
    ('AuthContext', 'Global React context exposing the authenticated user, token, and login/logout functions.'),
    ('Internationalization', 'All texts use the useTranslation() hook from react-i18next with JSON files per language in /locales.'),
]
for pattern, desc in patterns:
    add_paragraph(doc, f'● {pattern}:', bold=True)
    add_paragraph(doc, f'  {desc}')

page_break(doc)

# ═══════════════════════════════════════════════
# 10. API DOCUMENTATION
# ═══════════════════════════════════════════════
add_heading(doc, '10. Web Services Catalog and API Documentation', level=1)

add_heading(doc, '10.1 General API Information', level=2)
api_info = [
    ('Base URL (development)', 'http://localhost:8080/api'),
    ('Base URL (production)', 'https://invernadero-backend.onrender.com/api'),
    ('Swagger UI', 'http://localhost:8080/swagger-ui.html'),
    ('OpenAPI Specification', 'http://localhost:8080/v3/api-docs'),
    ('Format', 'JSON (application/json)'),
    ('Authentication', 'Bearer JWT in Authorization header'),
    ('API Version', '1.0.0'),
]
for key, value in api_info:
    p = doc.add_paragraph()
    p.add_run(f'{key}: ').bold = True
    p.add_run(value)

doc.add_paragraph()
add_heading(doc, '10.2 Authentication Endpoints', level=2)

auth_endpoints = [
    ('POST', '/api/auth/registro', 'Public', '{ nombre, apellido, email, password, rol }', '{ token, user }', 'Register a new local user'),
    ('POST', '/api/auth/login', 'Public', '{ email, password }', '{ token, user }', 'Login with local credentials'),
    ('POST', '/api/auth/google', 'Public', '{ idToken }', '{ token, user }', 'Authentication with Google OAuth 2.0'),
    ('GET', '/api/auth/me', 'JWT', '—', '{ user }', 'Get authenticated user profile'),
]

table_auth = doc.add_table(rows=len(auth_endpoints)+1, cols=6)
table_auth.style = 'Table Grid'
table_auth.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_auth, ['Method', 'Endpoint', 'Auth', 'Request Body', 'Response', 'Description'])
for i, row_data in enumerate(auth_endpoints):
    add_data_row(table_auth, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '10.3 CRUD Endpoints by Entity', level=2)
add_paragraph(doc,
    'All system entities expose the same standard set of REST operations. '
    'The following table summarizes the pattern applied to each entity:'
)

entities_ep = [
    ('Greenhouse', '/api/invernadero', 'Full CRUD'),
    ('Zone', '/api/zona', 'Full CRUD'),
    ('Crop', '/api/cultivo', 'Full CRUD'),
    ('Planting', '/api/siembra', 'CRUD + status change'),
    ('Sensor', '/api/sensor', 'Full CRUD'),
    ('Sensor Reading', '/api/lectura_sensor', 'CRUD + filter by sensor/zone'),
    ('Irrigation', '/api/riego', 'Full CRUD'),
    ('Alert', '/api/alerta', 'CRUD + mark as resolved'),
    ('Supply', '/api/insumo', 'CRUD + stock control'),
    ('Supply Application', '/api/aplicacion_insumo', 'Full CRUD'),
    ('Harvest', '/api/cosecha', 'Full CRUD'),
]

table_ent_ep = doc.add_table(rows=len(entities_ep)+1, cols=3)
table_ent_ep.style = 'Table Grid'
table_ent_ep.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_ent_ep, ['Entity', 'Base Path', 'Operations'])
for i, row_data in enumerate(entities_ep):
    add_data_row(table_ent_ep, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '10.4 Standard Endpoint Pattern', level=2)
add_code_block(doc, 'GET    /api/{entity}         → List all (paginated)')
add_code_block(doc, 'GET    /api/{entity}/{id}    → Find by ID')
add_code_block(doc, 'POST   /api/{entity}         → Create new record')
add_code_block(doc, 'PUT    /api/{entity}/{id}    → Full update of record')
add_code_block(doc, 'DELETE /api/{entity}/{id}    → Delete record')

doc.add_paragraph()
add_heading(doc, '10.5 HTTP Response Codes', level=2)
http_codes = [
    ('200 OK', 'Successful GET or PUT operation'),
    ('201 Created', 'Successful POST operation, resource created'),
    ('204 No Content', 'Successful DELETE operation'),
    ('400 Bad Request', 'Invalid input data (validation failed)'),
    ('401 Unauthorized', 'JWT token missing, expired, or invalid'),
    ('403 Forbidden', 'Authenticated user lacks permission for the operation'),
    ('404 Not Found', 'Resource not found for the provided ID'),
    ('409 Conflict', 'Conflict (e.g., duplicate email on registration)'),
    ('500 Internal Server Error', 'Unhandled server-side error'),
]
table_http = doc.add_table(rows=len(http_codes)+1, cols=2)
table_http.style = 'Table Grid'
table_http.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_http, ['Code', 'Meaning'])
for i, row_data in enumerate(http_codes):
    add_data_row(table_http, i+1, row_data, alternate=(i % 2 == 0))

page_break(doc)

# ═══════════════════════════════════════════════
# 11. TESTING
# ═══════════════════════════════════════════════
add_heading(doc, '11. Testing (Unit, Functional and Integration)', level=1)

add_heading(doc, '11.1 Testing Strategy', level=2)
add_paragraph(doc,
    'The project implements a three-layer testing strategy covering the backend, frontend, '
    'and Python automation scripts. Tests run automatically in the GitHub Actions CI/CD '
    'pipeline on every push and pull request.'
)

add_heading(doc, '11.2 Backend Tests – JUnit 5 (Spring Boot Test)', level=2)
add_paragraph(doc, 'Tools: JUnit 5, Spring Boot Test, Spring Security Test, H2 in-memory database, MockMvc.', italic=True)
add_paragraph(doc, 'Command: mvn test -Dspring.profiles.active=test', bold=True)

junit_tests = [
    ('AuthControllerTest', 'testRegistroExitoso()', 'Integration', 'Valid user registration returns 201 + JWT token'),
    ('AuthControllerTest', 'testRegistroEmailDuplicado()', 'Integration', 'Duplicate email returns 409 Conflict'),
    ('AuthControllerTest', 'testLoginExitoso()', 'Integration', 'Login with valid credentials returns 200 + token'),
    ('AuthControllerTest', 'testLoginCredencialesInvalidas()', 'Integration', 'Incorrect credentials return 401 Unauthorized'),
    ('AuthControllerTest', 'testRegistroConDatosInvalidos()', 'Unit', 'Empty required fields return 400 Bad Request'),
    ('InvernaderoControllerTest', 'testCrearInvernadero()', 'Integration', 'POST /api/invernadero returns 201 and the created object'),
    ('InvernaderoControllerTest', 'testListarInvernaderos()', 'Integration', 'GET /api/invernadero returns paginated list with 200'),
    ('InvernaderoControllerTest', 'testActualizarInvernadero()', 'Integration', 'PUT /api/invernadero/{id} updates and returns 200'),
    ('InvernaderoControllerTest', 'testEliminarInvernadero()', 'Integration', 'DELETE /api/invernadero/{id} returns 204 No Content'),
    ('InvernaderoControllerTest', 'testAccesoSinAutenticacion()', 'Security', 'Without JWT token returns 401 Unauthorized'),
]

table_junit = doc.add_table(rows=len(junit_tests)+1, cols=4)
table_junit.style = 'Table Grid'
table_junit.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_junit, ['Test Class', 'Method', 'Type', 'Description / Expected Result'])
for i, row_data in enumerate(junit_tests):
    add_data_row(table_junit, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '11.3 Frontend Tests – Selenium WebDriver (Python)', level=2)
add_paragraph(doc, 'Tools: Selenium 4, pytest, webdriver-manager, headless Chrome.', italic=True)
add_paragraph(doc, 'Command: pytest Frontend/tests/test_selenium.py -v', bold=True)
add_paragraph(doc, 'Prerequisites: Backend running on :8080, Frontend on :5173.', italic=True)

selenium_tests = [
    ('TestAuthentication', 'test_login_page_loads()', 'E2E', '/login page loads correctly with the form visible'),
    ('TestAuthentication', 'test_user_registration()', 'E2E', 'Full registration flow: form → submit → redirect to dashboard'),
    ('TestAuthentication', 'test_login_invalid_credentials()', 'E2E', 'Login with wrong credentials shows an error message'),
    ('TestNavigation', 'test_navigate_to_greenhouses()', 'E2E', 'Sidebar navigates correctly to /invernadero'),
    ('TestNavigation', 'test_navigate_to_sensors()', 'E2E', 'Sidebar navigates correctly to /sensor'),
    ('TestCRUDGreenhouse', 'test_create_greenhouse()', 'Functional', 'Create modal: fill form → save → appears in table'),
    ('TestCRUDGreenhouse', 'test_list_greenhouses()', 'Functional', 'Greenhouse table shows records with correct columns'),
    ('TestInternationalization', 'test_language_switch_es_en()', 'Functional', 'Language selector changes all visible texts to English'),
]

table_sel = doc.add_table(rows=len(selenium_tests)+1, cols=4)
table_sel.style = 'Table Grid'
table_sel.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_sel, ['Test Class', 'Method', 'Type', 'Description / Expected Result'])
for i, row_data in enumerate(selenium_tests):
    add_data_row(table_sel, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '11.4 Python Tests – pytest (Generator Scripts)', level=2)
add_paragraph(doc, 'Tools: pytest, coverage.py.', italic=True)
add_paragraph(doc, 'Command: pytest tests/test_generadores.py -v --cov=. --cov-report=html', bold=True)

pytest_tests = [
    ('TestSQLGeneration', 'test_generate_sql_creates_tables()', 'Unit', 'Generated SQL contains CREATE TABLE statements for all 12 entities'),
    ('TestSQLGeneration', 'test_generate_sql_includes_primary_keys()', 'Unit', 'Each table has its PRIMARY KEY defined'),
    ('TestSQLGeneration', 'test_generate_sql_includes_foreign_keys()', 'Unit', 'FK relationships are correctly defined'),
    ('TestSQLGeneration', 'test_generate_sql_includes_constraints()', 'Unit', 'NOT NULL, UNIQUE, and DEFAULT constraints are present'),
    ('TestDictionaryGeneration', 'test_dictionary_includes_entities()', 'Unit', 'The TXT file contains all 12 entities'),
    ('TestDictionaryGeneration', 'test_dictionary_includes_fields()', 'Unit', 'Each entity lists all its fields'),
    ('TestRelationshipGeneration', 'test_relationships_include_foreign_keys()', 'Unit', 'The textual ERD shows all FK relationships'),
    ('TestModelValidation', 'test_json_model_is_valid()', 'Unit', 'The source JSON is parseable and valid'),
    ('TestModelValidation', 'test_all_entities_have_id()', 'Unit', 'All entities have a defined PK field'),
    ('TestModelValidation', 'test_relationships_are_valid()', 'Integrity', 'All FKs reference entities that exist in the model'),
    ('TestGeneratedFiles', 'test_sql_file_exists()', 'Integration', 'The .sql file exists after running the generator'),
    ('TestGeneratedFiles', 'test_dictionary_pdf_exists()', 'Integration', 'The PDF dictionary is generated correctly'),
]

table_pytest = doc.add_table(rows=len(pytest_tests)+1, cols=4)
table_pytest.style = 'Table Grid'
table_pytest.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_pytest, ['Test Class', 'Method', 'Type', 'Description / Expected Result'])
for i, row_data in enumerate(pytest_tests):
    add_data_row(table_pytest, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '11.5 Coverage Goals', level=2)
coverage_data = [
    ('Backend (JUnit + JaCoCo)', '> 80%', 'mvn test jacoco:report → target/site/jacoco/index.html'),
    ('Frontend (Selenium)', 'Critical E2E flows', 'Login, main CRUD, internationalization'),
    ('Python (pytest + coverage)', '100% generator functions', 'pytest --cov=. → htmlcov/index.html'),
]
table_cov = doc.add_table(rows=len(coverage_data)+1, cols=3)
table_cov.style = 'Table Grid'
table_cov.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_cov, ['Layer', 'Target', 'Verification'])
for i, row_data in enumerate(coverage_data):
    add_data_row(table_cov, i+1, row_data, alternate=(i % 2 == 0))

page_break(doc)

# ═══════════════════════════════════════════════
# 12. ARCHITECTURE
# ═══════════════════════════════════════════════
add_heading(doc, '12. Proposed Architecture Diagram', level=1)

add_heading(doc, '12.1 Architecture Description', level=2)
add_paragraph(doc,
    'The system follows a decoupled three-tier architecture (Three-Tier Architecture) where '
    'the frontend and backend communicate exclusively through the REST API. '
    'Deployment is entirely cloud-based using free-tier services.'
)

add_heading(doc, '12.2 Architecture Components', level=2)

arch_data = [
    ('Web Client (React 19)', 'Vercel (Global CDN)', 'SPA with React + Vite, consumes the REST API via Axios, Google OAuth authentication.'),
    ('REST API (Spring Boot 3.2)', 'Render.com (Free Tier)', 'Java 17 backend with Spring Security, JPA, JWT. Exposes 60+ REST endpoints documented with Swagger.'),
    ('Database', 'Supabase (PostgreSQL)', 'Relational database with 12 tables, indexes, and referential integrity constraints.'),
    ('OAuth Authentication', 'Google Cloud (APIs)', 'Validation of Google ID tokens for passwordless authentication.'),
    ('CI/CD Pipeline', 'GitHub Actions', 'Automated pipeline: pytest → JUnit → Selenium → Build → Deploy → Taiga notification.'),
    ('Project Management', 'Taiga (Cloud)', 'User stories, acceptance criteria, tasks auto-created from CI/CD on failure.'),
    ('Automation Scripts', 'Local / CI', 'Python: generates SQL, dictionary, Java backend and React frontend from base_datos.json.'),
]

table_arch = doc.add_table(rows=len(arch_data)+1, cols=3)
table_arch.style = 'Table Grid'
table_arch.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_arch, ['Component', 'Platform', 'Responsibilities'])
for i, row_data in enumerate(arch_data):
    add_data_row(table_arch, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '12.3 Architecture Diagram (Textual)', level=2)

arch_diagram = """
┌─────────────────────────────────────────────────────────────────────┐
│                      END USER (Browser)                              │
└─────────────────────────────┬───────────────────────────────────────┘
                              │ HTTPS
                              ▼
┌─────────────────────────────────────────────────────────────────────┐
│                FRONTEND – React 19 + Vite 6                         │
│                 Vercel (CDN – Global Distribution)                   │
│  ┌─────────────┐  ┌──────────────┐  ┌─────────────────────────┐   │
│  │  react-     │  │  react-      │  │  @react-oauth/google    │   │
│  │  router-dom │  │  i18next     │  │  (OAuth 2.0 Button)     │   │
│  └─────────────┘  └──────────────┘  └─────────────────────────┘   │
└─────────────────────────────┬───────────────────────────────────────┘
                              │ REST (JSON) + Bearer JWT
                              ▼
┌─────────────────────────────────────────────────────────────────────┐
│              BACKEND – Spring Boot 3.2 / Java 17                    │
│                  Render.com (Free Web Service)                       │
│  ┌────────────────┐  ┌─────────────────┐  ┌──────────────────────┐ │
│  │  Controllers   │  │  Spring Security│  │  SpringDoc/Swagger   │ │
│  │  (60+ endpoints│  │  (JWT + OAuth2) │  │  (/swagger-ui.html) │ │
│  └───────┬────────┘  └─────────────────┘  └──────────────────────┘ │
│          │                                                           │
│  ┌───────▼────────┐  ┌─────────────────┐                           │
│  │  Services      │  │  JPA Repositories│                           │
│  │  (Business    │◄─┤  (Spring Data)   │                           │
│  │   Logic)      │  └────────┬─────────┘                           │
│  └────────────────┘           │                                      │
└─────────────────────────────┬─┼─────────────────────────────────────┘
                              │ │ JDBC
                              ▼ ▼
┌─────────────────────────────────────────────────────────────────────┐
│              DATABASE – PostgreSQL 15                               │
│                   Supabase (Free Tier)                               │
│              12 Tables | 17 FK Relationships                        │
└─────────────────────────────────────────────────────────────────────┘
                              │
                              │ (ID Token)
                              ▼
┌─────────────────────────────────────────────────────────────────────┐
│              GOOGLE CLOUD – OAuth 2.0 APIs                          │
│              Google identity token validation                       │
└─────────────────────────────────────────────────────────────────────┘

CI/CD PIPELINE (GitHub Actions)
────────────────────────────────
Push/PR → pytest → JUnit → Selenium → Build JAR → Build Frontend
         ↓ (on failure)
       Auto-creates task in TAIGA → Team notification
         ↓ (on full pass)
       Deploy Backend  → Render.com
       Deploy Frontend → Vercel
"""

for line in arch_diagram.split('\n'):
    add_code_block(doc, line)

page_break(doc)

# ═══════════════════════════════════════════════
# 13. RESULTS AND DISCUSSION
# ═══════════════════════════════════════════════
add_heading(doc, '13. Results and Discussion', level=1)

add_heading(doc, '13.1 Results per Rubric Item', level=2)

results_data = [
    ('Item 1', 'JSON Model, DB, ER, Data Dictionary', '✅ Complete',
     'All artifacts were automatically generated from Python: base_datos_invernadero.json (12 entities), '
     'crear_db_postgresql.sql (15KB), diccionario_datos.txt (23KB), diccionario_datos.pdf (25KB), '
     'and entidades_relaciones.txt. The script generar_base_de_datos.py acts as the single source of truth.'),
    ('Item 2', 'Spring Boot Backend', '✅ Complete',
     'REST backend with Spring Boot 3.2 / Java 17 is operational. Layered architecture '
     '(controller–service–repository). 60+ endpoints documented with SpringDoc/Swagger. '
     'Automatically generated by generar_desde_plantilla.py.'),
    ('Item 3', 'React Frontend', '✅ Complete',
     'SPA with React 19 + Vite 6 + Tailwind CSS 4. 13 CRUD modules auto-generated. '
     'Dashboard, protected route system, AuthContext, and responsive Layout implemented.'),
    ('Item 4', 'JUnit + Selenium + pytest Tests', '✅ Complete',
     '11 JUnit tests (AuthController + InvernaderoController), 8 Selenium E2E tests (authentication, '
     'navigation, CRUD, i18n), and 12 pytest tests for generator scripts. Backend coverage > 80%.'),
    ('Item 5', 'GitHub Actions CI/CD', '✅ Complete',
     'ci-cd.yml pipeline with sequential jobs: test-python → test-backend → test-frontend → deploy. '
     'Automatic Taiga integration: creates task on failure. Auto-deploy to Render and Vercel.'),
    ('Item 6', 'Google OAuth + Spring Security', '✅ Complete',
     'Dual authentication: local (BCrypt + JWT) and Google OAuth 2.0 (Google ID Token validated server-side). '
     'Roles and permissions implemented. SecurityConfig with custom JWT filter.'),
    ('Item 7', 'Internationalization', '✅ Complete',
     'Frontend: react-i18next with JSON files in /locales/es and /locales/en. '
     'Backend: error messages and logs internationalized. Language switch without page reload.'),
    ('Item 8', 'API Documentation / Javadoc', '✅ Complete',
     'All endpoints annotated with @Operation, @ApiResponse, @Tag from SpringDoc. '
     'Swagger UI available at /swagger-ui.html. Javadoc headers and comments in all files.'),
    ('Item 9', 'Taiga Integration', '✅ Complete',
     'actualizarToken_taiga.py for token management. CI/CD creates tasks automatically when tests fail. '
     'User story acceptance criteria linked to the pipeline.'),
]

for item_data in results_data:
    add_heading(doc, f'{item_data[0]}: {item_data[1]} — {item_data[2]}', level=3)
    add_paragraph(doc, item_data[3])

add_heading(doc, '13.2 Notable Technical Achievements', level=2)
achievements = [
    'Automated code generation: a single JSON as the source of truth generates SQL, the data dictionary, Java backend, and React frontend — eliminating duplication and guaranteeing consistency.',
    'Complete CI/CD pipeline with three testing layers (Python → Java → Selenium) and automatic task creation in Taiga on failure, closing the quality cycle.',
    'Secure dual authentication: local login support with BCrypt + JWT and Google OAuth 2.0 login, with ID token validation on the server side.',
    'Decoupled architecture: frontend and backend are fully independent, deployable on different platforms, communicating only via REST API.',
    'Complete internationalization: real-time support for Spanish and English in the frontend and in backend error messages.',
]
for achievement in achievements:
    add_bullet(doc, achievement)

add_heading(doc, '13.3 Limitations and Future Work', level=2)
limitations = [
    ('Physical Sensors', 'The system currently simulates sensor readings. A real-world deployment would require integration with IoT devices (Arduino/Raspberry Pi) via MQTT or WebSockets.'),
    ('Real-Time Notifications', 'Alerts currently require the user to refresh the view. WebSocket or Server-Sent Events (SSE) could be implemented for real-time push notifications.'),
    ('Machine Learning', 'The alert system uses fixed thresholds. A more advanced approach could apply prediction models to anticipate adverse conditions before they occur.'),
    ('Sensor Scalability', 'For a large greenhouse with hundreds of sensors, the SENSOR_READING table would grow rapidly. Migrating to a time-series database (e.g., InfluxDB) is recommended for this entity.'),
    ('Reporting Module', 'There is currently no exportable reports module (PDF/Excel) with historical production analysis and performance by crop type.'),
]
for title, desc in limitations:
    add_paragraph(doc, f'● {title}:', bold=True)
    add_paragraph(doc, f'  {desc}')

add_heading(doc, '13.4 Conclusions', level=2)
add_paragraph(doc,
    'The Automated Greenhouse Management System satisfactorily meets all 9 items of the academic '
    'rubric, demonstrating the feasibility of building a complete web system with development '
    'process automation, test coverage across all layers, robust security, internationalization, '
    'and continuous deployment.'
)
add_paragraph(doc,
    'The adoption of the code generation pattern from a centralized JSON model was the most '
    'significant architectural decision of the project. It guaranteed consistency between the '
    'data model and the implementation, substantially reducing repetitive development effort '
    'and synchronization errors.'
)
add_paragraph(doc,
    'The selected technology stack (Spring Boot + React + PostgreSQL + GitHub Actions + '
    'free cloud tools) proved to be a robust, scalable, and cost-free combination, suitable '
    'both for the academic context and for a real production scenario at small or medium scale.'
)

# ═══════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════
page_break(doc)
add_heading(doc, 'References', level=1)
references = [
    'Spring Boot Reference Documentation v3.2. Pivotal / VMware. https://docs.spring.io/spring-boot/docs/3.2.x/reference/html/',
    'React Documentation v19. Meta Open Source. https://react.dev/',
    'SpringDoc OpenAPI Documentation v2.5. https://springdoc.org/',
    'JJWT (Java JWT) v0.12.5. https://github.com/jwtk/jjwt',
    'Google OAuth 2.0 for Web Server Applications. Google Developers. https://developers.google.com/identity/protocols/oauth2',
    'Selenium WebDriver Documentation. https://www.selenium.dev/documentation/',
    'pytest Documentation v8. https://docs.pytest.org/',
    'JUnit 5 User Guide. https://junit.org/junit5/docs/current/user-guide/',
    'Priva Greenhouse Manager. https://www.priva.com/solutions/greenhouse-management',
    'Taiga Project Management. https://taiga.io/',
    'Supabase PostgreSQL. https://supabase.com/',
    'Render.com Deployment Documentation. https://render.com/docs/',
    'Vercel Deployment Documentation. https://vercel.com/docs',
    'GitHub Actions Documentation. https://docs.github.com/en/actions',
    'react-i18next Documentation. https://react.i18next.com/',
]
for ref in references:
    add_bullet(doc, ref)

# ═══════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════
output_path = r'C:\Users\Usuario\Desktop\Invernadero automatizado\Documentation_Automated_Greenhouse.docx'
doc.save(output_path)
print(f'✅ Document generated successfully: {output_path}')
print(f'📄 Sections included:')
for num, section in toc_items:
    print(f'   {num} {section}')
