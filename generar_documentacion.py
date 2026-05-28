"""
generar_documentacion.py
Genera la documentación completa del proyecto Invernadero Automatizado en formato Word (.docx).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────
# Helpers de estilo
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '2E7D32')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_header_row(table, headers, bg='2E7D32', fg='FFFFFF'):
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(fg)
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, bg)
        set_cell_border(cell)

def add_data_row(table, row_index, values, alternate=False):
    row = table.rows[row_index]
    bg = 'E8F5E9' if alternate else 'FFFFFF'
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(str(val))
        run.font.size = Pt(8.5)
        set_cell_bg(cell, bg)
        set_cell_border(cell)

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0] if heading.runs else heading.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
        run.font.size = Pt(13)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x38, 0x8E, 0x3C)
        run.font.size = Pt(11)
    return heading

def add_paragraph(doc, text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x01, 0x57, 0x9B)

def page_break(doc):
    doc.add_page_break()


# ─────────────────────────────────────────────
# DOCUMENTO
# ─────────────────────────────────────────────

doc = Document()

# Márgenes
sections = doc.sections
for section in sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# Estilos de fuente base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ═══════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SISTEMA DE GESTIÓN DE\nINVERNADERO AUTOMATIZADO')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('Documentación Técnica del Proyecto')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f'Versión: 1.0.0\n').font.size = Pt(11)
info.add_run(f'Fecha: {datetime.date.today().strftime("%d de %B de %Y")}\n').font.size = Pt(11)
info.add_run('Materia: Fábrica de Software').font.size = Pt(11)

page_break(doc)

# ═══════════════════════════════════════════════
# ÍNDICE (manual)
# ═══════════════════════════════════════════════
add_heading(doc, 'Tabla de Contenido', level=1)

toc_items = [
    ('1.', 'Introducción'),
    ('2.', 'Problema'),
    ('3.', 'Objetivos'),
    ('4.', 'Estado del Arte – Fundamentación y Trabajos Previos'),
    ('5.', 'Requerimientos'),
    ('6.', 'Casos de Uso e Historias de Usuario'),
    ('7.', 'Diccionario de Datos y Modelo Entidad-Relación'),
    ('8.', 'Diagrama de Clases'),
    ('9.', 'Diseño de Interfaz Gráfica de Usuario'),
    ('10.', 'Catálogo de Servicios Web y Documentación de APIs'),
    ('11.', 'Pruebas (Unitarias, Funcionales e Integración)'),
    ('12.', 'Gráfica de Arquitectura del Modelo Propuesto'),
    ('13.', 'Resultados y Discusión'),
]

for num, item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'{num}  {item}')
    run.font.size = Pt(11)

page_break(doc)

# ═══════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ═══════════════════════════════════════════════
add_heading(doc, '1. Introducción', level=1)
add_paragraph(doc,
    'El presente documento describe el diseño, desarrollo e implementación del Sistema de Gestión '
    'de Invernadero Automatizado, un aplicativo web desarrollado como proyecto académico para la materia '
    'de Fábrica de Software. El sistema tiene como propósito centralizar y automatizar los procesos '
    'operativos críticos de un invernadero: monitoreo ambiental mediante sensores, gestión de siembras '
    'y cosechas, control de riego, administración de insumos y generación de alertas ante condiciones '
    'anómalas.'
)
add_paragraph(doc,
    'La solución está construida sobre una arquitectura de tres capas: un backend REST desarrollado '
    'con Spring Boot 3.2 y Java 17, un frontend SPA (Single Page Application) desarrollado con '
    'React 19 y Vite 6, y una base de datos relacional PostgreSQL. El proyecto también incorpora '
    'automatización mediante scripts Python para la generación de artefactos (modelo de base de '
    'datos, diccionario de datos, estructura del backend y del frontend), un pipeline CI/CD con '
    'GitHub Actions, autenticación con Google OAuth 2.0, internacionalización y pruebas en todas '
    'las capas del sistema.'
)
add_paragraph(doc,
    'Uno de los aspectos diferenciadores de este proyecto es el enfoque en la automatización del '
    'proceso de desarrollo: el modelo de datos se define una única vez en un archivo JSON y a partir '
    'de él se generan automáticamente el esquema SQL, el diccionario de datos, los controladores '
    'REST del backend y las páginas CRUD del frontend.'
)

page_break(doc)

# ═══════════════════════════════════════════════
# 2. PROBLEMA
# ═══════════════════════════════════════════════
add_heading(doc, '2. Problema', level=1)

add_heading(doc, '2.1 Descripción del Problema', level=2)
add_paragraph(doc,
    'Los invernaderos modernos enfrentan desafíos significativos en la gestión manual de sus '
    'procesos. Los operarios deben registrar lecturas de temperatura, humedad, pH y luminosidad '
    'de manera manual, lo que introduce errores humanos y retrasos en la detección de condiciones '
    'críticas que pueden comprometer los cultivos. Adicionalmente, la gestión del inventario de '
    'insumos, el registro de riegos y el seguimiento del ciclo de vida de las siembras se realizan '
    'frecuentemente en hojas de cálculo o en papel, lo que dificulta la trazabilidad y el análisis.'
)

add_heading(doc, '2.2 Causas Identificadas', level=2)
add_bullet(doc, 'Ausencia de un sistema centralizado de monitoreo de variables ambientales.')
add_bullet(doc, 'Procesos manuales propensos a error humano en el registro de operaciones.')
add_bullet(doc, 'Falta de alertas automáticas ante condiciones fuera de rango óptimo.')
add_bullet(doc, 'Dificultad para realizar trazabilidad de los ciclos de cultivo completos.')
add_bullet(doc, 'Falta de acceso multi-rol con niveles de permisos diferenciados.')
add_bullet(doc, 'Ausencia de reportes y visualizaciones del estado del invernadero en tiempo real.')

add_heading(doc, '2.3 Impacto del Problema', level=2)
add_paragraph(doc,
    'La falta de automatización genera pérdidas de cultivos por reacción tardía a condiciones '
    'ambientales adversas, desperdicio de insumos por falta de control de inventario y '
    'dificultad para identificar patrones y optimizar los procesos productivos del invernadero.'
)

page_break(doc)

# ═══════════════════════════════════════════════
# 3. OBJETIVOS
# ═══════════════════════════════════════════════
add_heading(doc, '3. Objetivos', level=1)

add_heading(doc, '3.1 Objetivo General', level=2)
add_paragraph(doc,
    'Desarrollar un sistema web de gestión integral para invernaderos que automatice el monitoreo '
    'de variables ambientales, el control de riego, la gestión de cultivos y la administración de '
    'insumos, integrando una arquitectura REST segura, internacionalizada, con despliegue continuo '
    'y cobertura de pruebas en todas las capas.'
)

add_heading(doc, '3.2 Objetivos Específicos', level=2)
especificos = [
    'Diseñar y generar automáticamente el modelo de base de datos relacional en PostgreSQL a partir de un modelo JSON fuente de verdad.',
    'Implementar un backend REST con Spring Boot 3.2 que exponga los servicios CRUD para las 12 entidades del sistema, con seguridad basada en JWT y autenticación OAuth 2.0 con Google.',
    'Desarrollar un frontend SPA con React 19 que consuma los servicios REST y soporte internacionalización (español e inglés).',
    'Configurar un pipeline CI/CD con GitHub Actions que ejecute pruebas automáticas (JUnit, Selenium, pytest) y genere tareas en Taiga ante fallos.',
    'Implementar un sistema de alertas que detecte condiciones anómalas en las lecturas de sensores y notifique a los usuarios correspondientes.',
    'Documentar la API REST mediante SpringDoc OpenAPI (Swagger UI) con anotaciones completas en todos los endpoints.',
    'Alcanzar una cobertura de pruebas superior al 80% en el backend y cubrir los flujos críticos del frontend con pruebas E2E (Selenium).',
    'Automatizar la generación de componentes del sistema (backend y frontend) mediante scripts Python.',
]
for obj in especificos:
    add_bullet(doc, obj)

page_break(doc)

# ═══════════════════════════════════════════════
# 4. ESTADO DEL ARTE
# ═══════════════════════════════════════════════
add_heading(doc, '4. Estado del Arte – Fundamentación y Trabajos Previos', level=1)

add_heading(doc, '4.1 Invernaderos Inteligentes e IoT', level=2)
add_paragraph(doc,
    'El concepto de invernadero inteligente (Smart Greenhouse) ha evolucionado significativamente '
    'en la última década impulsado por el Internet de las Cosas (IoT). Sistemas como Greenhouse '
    'Manager (Priva), iGrow y Argus Controls permiten la integración de sensores de temperatura, '
    'humedad relativa, CO₂ y pH con sistemas de control automático de ventilación, riego y '
    'fertilización. Sin embargo, estas soluciones comerciales presentan costos de licencia elevados '
    'y escasa flexibilidad para adaptarse a pequeñas y medianas unidades productivas.'
)

add_heading(doc, '4.2 Trabajos Académicos Previos', level=2)
trabajos = [
    ('Monitoreo IoT con Arduino y MQTT (2022)',
     'Sistemas de bajo costo basados en Arduino + Raspberry Pi con protocolo MQTT para transmisión '
     'de lecturas de sensores hacia plataformas como ThingsBoard. Validaron la viabilidad técnica '
     'del monitoreo remoto en tiempo real con costos de hardware inferiores a USD 100.'),
    ('Sistemas de Riego Automatizado con Machine Learning (2023)',
     'Investigaciones que combinan sensores de humedad de suelo con modelos de ML (Random Forest, '
     'Redes Neuronales) para predicción óptima del momento de riego, logrando reducciones de '
     'consumo hídrico del 30% comparado con sistemas de riego temporizados.'),
    ('Plataformas Web para Gestión Agrícola (2021-2024)',
     'Proyectos académicos que implementan arquitecturas REST + React para gestión de cultivos, '
     'destacando el uso de Spring Boot como backend por su ecosistema maduro, soporte para '
     'seguridad (Spring Security) y facilidad de documentación (SpringDoc/Swagger).'),
    ('Automatización con Generación de Código (2024)',
     'Proyectos de fábrica de software que implementan el patrón de generación de código a '
     'partir de modelos JSON, reduciendo el tiempo de desarrollo repetitivo y garantizando '
     'consistencia entre el modelo de datos y la implementación.'),
]
for titulo, desc in trabajos:
    add_paragraph(doc, f'● {titulo}', bold=True)
    add_paragraph(doc, f'  {desc}')

add_heading(doc, '4.3 Tecnologías Utilizadas y Justificación', level=2)
tech_rows = [
    ('Spring Boot 3.2 / Java 17', 'Framework maduro para REST APIs con ecosistema de seguridad, JPA y documentación integrados.'),
    ('React 19 + Vite 6', 'SPA con rendimiento óptimo, ecosistema robusto y soporte para i18n y OAuth.'),
    ('PostgreSQL', 'Base de datos relacional robusta, con soporte avanzado de tipos y transacciones ACID.'),
    ('Python 3.10+', 'Scripts de automatización para generación de artefactos y pruebas con pytest.'),
    ('GitHub Actions', 'CI/CD nativo de GitHub, gratuito para repositorios públicos/privados.'),
    ('Google OAuth 2.0', 'Autenticación delegada a proveedor de identidad confiable, sin almacenar contraseñas.'),
    ('JWT (JJWT 0.12.5)', 'Tokens stateless para autorización, escalables y compatibles con arquitecturas distribuidas.'),
    ('SpringDoc OpenAPI 2.5', 'Generación automática de documentación Swagger UI desde anotaciones Java.'),
    ('Selenium + WebDriver', 'Pruebas E2E del frontend en navegador real, estándar de la industria.'),
    ('Taiga', 'Plataforma ágil de gestión de proyectos integrada con el CI/CD para trazabilidad.'),
]
table_tech = doc.add_table(rows=len(tech_rows)+1, cols=2)
table_tech.style = 'Table Grid'
table_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_tech, ['Tecnología', 'Justificación'])
for i, row_data in enumerate(tech_rows):
    add_data_row(table_tech, i+1, row_data, alternate=(i % 2 == 0))

page_break(doc)

# ═══════════════════════════════════════════════
# 5. REQUERIMIENTOS
# ═══════════════════════════════════════════════
add_heading(doc, '5. Requerimientos', level=1)

add_heading(doc, '5.1 Requerimientos Funcionales', level=2)

rf_data = [
    ('RF-01', 'Gestión de Usuarios', 'El sistema debe permitir el registro, autenticación local y OAuth Google, edición y desactivación de usuarios con roles diferenciados (ADMINISTRADOR, TÉCNICO, OPERARIO, VISUALIZADOR).', 'Alta'),
    ('RF-02', 'Gestión de Invernaderos', 'El sistema debe permitir crear, listar, actualizar y eliminar invernaderos con sus atributos de nombre, ubicación, área y tipo de estructura.', 'Alta'),
    ('RF-03', 'Gestión de Zonas', 'El sistema debe permitir la definición de zonas internas dentro de cada invernadero.', 'Alta'),
    ('RF-04', 'Catálogo de Cultivos', 'El sistema debe permitir registrar y mantener el catálogo de especies con sus parámetros óptimos (temperatura, humedad, ciclo).', 'Alta'),
    ('RF-05', 'Registro de Siembras', 'El sistema debe registrar los eventos de siembra asociando zona, cultivo, usuario y seguimiento del estado del ciclo.', 'Alta'),
    ('RF-06', 'Gestión de Sensores', 'El sistema debe permitir registrar y monitorear los sensores instalados en cada zona, con tipos: TEMPERATURA, HUMEDAD, CO2, LUMINOSIDAD, PH, HUMEDAD_SUELO.', 'Alta'),
    ('RF-07', 'Lecturas de Sensores', 'El sistema debe almacenar la serie temporal de lecturas de cada sensor e identificar aquellas que superen umbrales definidos.', 'Alta'),
    ('RF-08', 'Control de Riego', 'El sistema debe registrar eventos de riego (automático o manual) con duración y volumen.', 'Media'),
    ('RF-09', 'Sistema de Alertas', 'El sistema debe generar y notificar alertas clasificadas por nivel (INFORMATIVA, ADVERTENCIA, CRÍTICA) ante condiciones anómalas.', 'Alta'),
    ('RF-10', 'Gestión de Insumos', 'El sistema debe administrar el inventario de insumos con control de stock mínimo y alertas de reabastecimiento.', 'Media'),
    ('RF-11', 'Aplicación de Insumos', 'El sistema debe registrar las aplicaciones de insumos a siembras o zonas con trazabilidad completa.', 'Media'),
    ('RF-12', 'Registro de Cosechas', 'El sistema debe registrar la producción cosechada por siembra con clasificación de calidad.', 'Alta'),
    ('RF-13', 'Dashboard', 'El sistema debe presentar un panel de control con métricas clave del estado del invernadero.', 'Alta'),
    ('RF-14', 'Internacionalización', 'El sistema debe soportar la interfaz en español e inglés con posibilidad de cambio en tiempo de ejecución.', 'Media'),
    ('RF-15', 'Integración Taiga', 'El sistema CI/CD debe crear tareas en Taiga automáticamente cuando fallen los tests.', 'Media'),
    ('RF-16', 'Documentación API', 'Todos los endpoints REST deben estar documentados y accesibles mediante Swagger UI.', 'Alta'),
]

table_rf = doc.add_table(rows=len(rf_data)+1, cols=4)
table_rf.style = 'Table Grid'
table_rf.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_rf, ['ID', 'Nombre', 'Descripción', 'Prioridad'])
for i, row_data in enumerate(rf_data):
    add_data_row(table_rf, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '5.2 Requerimientos No Funcionales', level=2)

rnf_data = [
    ('RNF-01', 'Seguridad', 'Autenticación JWT con tokens de corta duración y validación de Google ID Token. Contraseñas almacenadas con BCrypt.'),
    ('RNF-02', 'Rendimiento', 'Las consultas de listas deben responder en menos de 500ms bajo carga normal. Paginación disponible en todos los endpoints de lista.'),
    ('RNF-03', 'Disponibilidad', 'El sistema debe garantizar un uptime mínimo del 99% en el ambiente de producción (Render.com).'),
    ('RNF-04', 'Escalabilidad', 'La arquitectura stateless (JWT) permite escalar horizontalmente sin configuración adicional.'),
    ('RNF-05', 'Mantenibilidad', 'Código documentado con Javadoc/docstrings. Cobertura de pruebas > 80%. Arquitectura en capas (controller → service → repository).'),
    ('RNF-06', 'Portabilidad', 'El backend corre en cualquier entorno con Java 17+. El frontend es un bundle estático desplegable en cualquier CDN.'),
    ('RNF-07', 'Usabilidad', 'Interfaz responsiva con soporte multilingüe y retroalimentación visual en todas las operaciones CRUD.'),
    ('RNF-08', 'Trazabilidad', 'Todos los eventos críticos (siembra, cosecha, alertas, aplicaciones) quedan registrados con timestamp y usuario responsable.'),
]
table_rnf = doc.add_table(rows=len(rnf_data)+1, cols=3)
table_rnf.style = 'Table Grid'
table_rnf.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_rnf, ['ID', 'Nombre', 'Descripción'])
for i, row_data in enumerate(rnf_data):
    add_data_row(table_rnf, i+1, row_data, alternate=(i % 2 == 0))

page_break(doc)

# ═══════════════════════════════════════════════
# 6. CASOS DE USO E HISTORIAS DE USUARIO
# ═══════════════════════════════════════════════
add_heading(doc, '6. Casos de Uso e Historias de Usuario', level=1)

add_heading(doc, '6.1 Actores del Sistema', level=2)
actores = [
    ('Administrador', 'Acceso total al sistema. Gestiona usuarios, invernaderos, configuraciones y visualiza todos los reportes.'),
    ('Técnico', 'Gestiona sensores, zonas, riegos y genera alertas. Puede registrar siembras y cosechas.'),
    ('Operario', 'Registra lecturas manuales, aplica insumos y ejecuta riegos. Visibilidad limitada a su zona asignada.'),
    ('Visualizador', 'Acceso de solo lectura a dashboards y reportes.'),
    ('Sistema (Automático)', 'Genera alertas por umbrales y crea tareas en Taiga por fallos de CI/CD.'),
]
for actor, desc in actores:
    add_paragraph(doc, f'● {actor}:', bold=True)
    add_paragraph(doc, f'  {desc}')

add_heading(doc, '6.2 Casos de Uso Principales', level=2)

cu_data = [
    ('CU-01', 'Autenticar Usuario', 'Operario, Técnico, Admin', 'Usuario con credenciales válidas o cuenta Google', 'Usuario autenticado con token JWT'),
    ('CU-02', 'Registrar Invernadero', 'Administrador', 'Datos del invernadero: nombre, ubicación, área, tipo', 'Invernadero registrado en base de datos'),
    ('CU-03', 'Registrar Lectura Sensor', 'Técnico, Sistema', 'Sensor activo y valor medido', 'Lectura almacenada, alerta generada si aplica'),
    ('CU-04', 'Gestionar Siembra', 'Técnico, Operario', 'Zona, cultivo y datos de siembra', 'Ciclo de siembra iniciado con trazabilidad'),
    ('CU-05', 'Ejecutar Riego', 'Operario, Sistema', 'Zona y parámetros de riego', 'Evento de riego registrado'),
    ('CU-06', 'Gestionar Alertas', 'Técnico, Administrador', 'Lista de alertas activas', 'Alertas resueltas y con fecha de resolución'),
    ('CU-07', 'Registrar Cosecha', 'Técnico, Operario', 'Siembra completada y datos de producción', 'Cosecha registrada con cantidad y calidad'),
    ('CU-08', 'Consultar Dashboard', 'Todos los roles', 'Sesión activa', 'Vista de métricas y estado general del invernadero'),
    ('CU-09', 'Gestionar Insumos', 'Técnico, Administrador', 'Datos del insumo y stock', 'Inventario actualizado con alertas de stock mínimo'),
    ('CU-10', 'Cambiar Idioma', 'Todos los roles', 'Selección de idioma (ES/EN)', 'Interfaz traducida en tiempo real'),
]

table_cu = doc.add_table(rows=len(cu_data)+1, cols=5)
table_cu.style = 'Table Grid'
table_cu.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_cu, ['ID', 'Caso de Uso', 'Actor(es)', 'Precondición', 'Postcondición'])
for i, row_data in enumerate(cu_data):
    add_data_row(table_cu, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '6.3 Historias de Usuario', level=2)

hu_data = [
    ('HU-01', 'Autenticación', 'Como usuario del sistema, quiero iniciar sesión con mi cuenta de Google para no tener que recordar otra contraseña.',
     '● Debe mostrar botón "Iniciar sesión con Google"\n● Al autenticar, redirige al Dashboard\n● El token JWT expira en 24h\n● Si el email no existe, se crea la cuenta automáticamente'),
    ('HU-02', 'Monitoreo de Sensores', 'Como técnico, quiero ver las lecturas en tiempo real de todos los sensores para detectar condiciones fuera de rango.',
     '● Lista todas las lecturas con sensor, valor y timestamp\n● Resalta en rojo las lecturas que generaron alerta\n● Permite filtrar por zona y tipo de sensor'),
    ('HU-03', 'Alertas Críticas', 'Como administrador, quiero recibir alertas clasificadas por nivel de gravedad para priorizar mi respuesta.',
     '● Las alertas se clasifican en INFORMATIVA, ADVERTENCIA y CRÍTICA\n● Se puede marcar una alerta como resuelta\n● Se registra la fecha y hora de resolución'),
    ('HU-04', 'Ciclo de Siembra', 'Como operario, quiero registrar una siembra nueva indicando la zona, el cultivo y la cantidad de plantas para llevar trazabilidad del ciclo.',
     '● Los campos obligatorios son: zona, cultivo, fecha siembra, cantidad de plantas\n● El estado inicial es EN_CRECIMIENTO\n● Se calcula y muestra la fecha estimada de cosecha'),
    ('HU-05', 'Inventario de Insumos', 'Como técnico, quiero ver una alerta cuando el stock de un insumo esté por debajo del mínimo para solicitar reabastecimiento a tiempo.',
     '● La tabla de insumos resalta en rojo los que están bajo el stock mínimo\n● Se puede registrar una nueva entrada al inventario\n● El historial de aplicaciones reduce el stock automáticamente'),
    ('HU-06', 'Internacionalización', 'Como usuario, quiero poder cambiar el idioma de la interfaz entre español e inglés en cualquier momento.',
     '● El cambio de idioma es inmediato sin recarga de página\n● La preferencia se guarda en localStorage\n● Todos los textos, mensajes de error y etiquetas están traducidos'),
    ('HU-07', 'Registro de Cosecha', 'Como técnico, quiero registrar la producción cosechada indicando la cantidad en kg y la calidad para tener trazabilidad completa del ciclo.',
     '● Se asocia a una siembra existente en estado EN_CRECIMIENTO o COSECHADO\n● La calidad puede ser PREMIUM, ESTÁNDAR, SEGUNDA o DESCARTE\n● Al registrar la cosecha, la siembra cambia a estado COSECHADO'),
    ('HU-08', 'Dashboard Ejecutivo', 'Como administrador, quiero ver un resumen del estado del invernadero en el dashboard para tomar decisiones informadas rápidamente.',
     '● Muestra total de invernaderos, zonas activas, alertas pendientes y siembras en curso\n● Gráficas de últimas lecturas de sensores por zona\n● Acceso rápido a las secciones más utilizadas'),
]

for hu in hu_data:
    add_heading(doc, f'{hu[0]}: {hu[1]}', level=3)
    add_paragraph(doc, hu[2], italic=True)
    add_paragraph(doc, 'Criterios de Aceptación:', bold=True)
    for line in hu[3].split('\n'):
        if line.strip():
            add_bullet(doc, line.strip().lstrip('●').strip())
    doc.add_paragraph()

page_break(doc)

# ═══════════════════════════════════════════════
# 7. DICCIONARIO DE DATOS Y MER
# ═══════════════════════════════════════════════
add_heading(doc, '7. Diccionario de Datos y Modelo Entidad-Relación', level=1)

add_heading(doc, '7.1 Resumen del Modelo', level=2)
add_paragraph(doc,
    'El modelo de datos del Sistema de Gestión de Invernadero Automatizado está compuesto por '
    '12 entidades que capturan la totalidad de los procesos operativos del invernadero. '
    'El modelo fue generado automáticamente a partir del archivo base_datos_invernadero.json '
    'mediante el script Python generar_base_de_datos.py.'
)

entidades_resumen = [
    ('USUARIO', 9, 'Usuarios del sistema con roles y permisos de acceso'),
    ('INVERNADERO', 8, 'Información general de cada instalación física de invernadero'),
    ('ZONA', 5, 'Secciones o áreas internas dentro de un invernadero'),
    ('CULTIVO', 10, 'Catálogo maestro de especies y variedades cultivadas'),
    ('SIEMBRA', 9, 'Registro de cada evento de siembra en una zona'),
    ('SENSOR', 7, 'Dispositivos de medición instalados en las zonas'),
    ('LECTURA_SENSOR', 5, 'Serie temporal de mediciones registradas por cada sensor'),
    ('RIEGO', 8, 'Eventos de riego automáticos o manuales por zona'),
    ('ALERTA', 9, 'Notificaciones por condiciones anómalas detectadas'),
    ('INSUMO', 7, 'Catálogo de insumos con control de inventario'),
    ('APLICACION_INSUMO', 9, 'Historial de aplicaciones de insumos a siembras o zonas'),
    ('COSECHA', 7, 'Registro de producción cosechada por evento de siembra'),
]

table_ent = doc.add_table(rows=len(entidades_resumen)+1, cols=3)
table_ent.style = 'Table Grid'
table_ent.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_ent, ['Entidad', 'N° Campos', 'Descripción'])
for i, row_data in enumerate(entidades_resumen):
    add_data_row(table_ent, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '7.2 Diccionario de Datos por Entidad', level=2)

# Datos de cada entidad extraídos del diccionario_datos.txt
entidades_dict = {
    'USUARIO': [
        ('id_usuario', 'INT', 'NO', 'SI', 'NO', 'NO', '—', 'Identificador único del usuario'),
        ('nombre', 'VARCHAR(100)', 'NO', 'NO', 'NO', 'NO', '—', 'Nombre del usuario'),
        ('apellido', 'VARCHAR(100)', 'NO', 'NO', 'NO', 'NO', '—', 'Apellido del usuario'),
        ('email', 'VARCHAR(150)', 'NO', 'NO', 'NO', 'SI', '—', 'Correo electrónico único para autenticación OAuth'),
        ('password_hash', 'VARCHAR(255)', 'SI', 'NO', 'NO', 'NO', '—', 'Hash de contraseña (nulo si usa solo OAuth)'),
        ('rol', 'VARCHAR(30)', 'NO', 'NO', 'NO', 'NO', 'OPERARIO', 'Rol: ADMINISTRADOR|TÉCNICO|OPERARIO|VISUALIZADOR'),
        ('fecha_registro', 'DATE', 'NO', 'NO', 'NO', 'NO', '—', 'Fecha de registro en el sistema'),
        ('activo', 'BOOLEAN', 'NO', 'NO', 'NO', 'NO', 'True', 'Indica si la cuenta está habilitada'),
        ('proveedor_oauth', 'VARCHAR(30)', 'SI', 'NO', 'NO', 'NO', '—', 'Proveedor OAuth: LOCAL|GOOGLE'),
    ],
    'INVERNADERO': [
        ('id_invernadero', 'INT', 'NO', 'SI', 'NO', 'NO', '—', 'Identificador único del invernadero'),
        ('nombre', 'VARCHAR(100)', 'NO', 'NO', 'NO', 'NO', '—', 'Nombre o código del invernadero'),
        ('ubicacion', 'VARCHAR(200)', 'NO', 'NO', 'NO', 'NO', '—', 'Descripción de la ubicación física'),
        ('area_m2', 'DECIMAL(10,2)', 'NO', 'NO', 'NO', 'NO', '—', 'Área total en metros cuadrados'),
        ('tipo_estructura', 'VARCHAR(50)', 'SI', 'NO', 'NO', 'NO', '—', 'Material: VIDRIO|POLICARBONATO|MALLA|PLÁSTICO'),
        ('responsable_id', 'INT', 'NO', 'NO', 'SI', 'NO', '—', 'FK → usuario.id_usuario'),
        ('fecha_creacion', 'DATE', 'NO', 'NO', 'NO', 'NO', '—', 'Fecha de registro en el sistema'),
        ('estado', 'VARCHAR(20)', 'NO', 'NO', 'NO', 'NO', 'ACTIVO', 'Estado: ACTIVO|INACTIVO|MANTENIMIENTO'),
    ],
    'SENSOR': [
        ('id_sensor', 'INT', 'NO', 'SI', 'NO', 'NO', '—', 'Identificador único del sensor'),
        ('id_zona', 'INT', 'NO', 'NO', 'SI', 'NO', '—', 'FK → zona.id_zona'),
        ('tipo_sensor', 'VARCHAR(50)', 'NO', 'NO', 'NO', 'NO', '—', 'TEMPERATURA|HUMEDAD|CO2|LUMINOSIDAD|PH|HUMEDAD_SUELO'),
        ('modelo', 'VARCHAR(100)', 'SI', 'NO', 'NO', 'NO', '—', 'Modelo o referencia comercial del sensor'),
        ('unidad_medida', 'VARCHAR(20)', 'NO', 'NO', 'NO', 'NO', '—', 'Unidad de medida (°C, %, ppm, lux)'),
        ('fecha_instalacion', 'DATE', 'SI', 'NO', 'NO', 'NO', '—', 'Fecha de instalación'),
        ('estado', 'VARCHAR(20)', 'NO', 'NO', 'NO', 'NO', 'ACTIVO', 'Estado: ACTIVO|INACTIVO|FALLA'),
    ],
    'ALERTA': [
        ('id_alerta', 'INT', 'NO', 'SI', 'NO', 'NO', '—', 'Identificador único de la alerta'),
        ('id_sensor', 'INT', 'SI', 'NO', 'SI', 'NO', '—', 'FK → sensor.id_sensor (nulo si alerta manual)'),
        ('id_zona', 'INT', 'SI', 'NO', 'SI', 'NO', '—', 'FK → zona.id_zona'),
        ('tipo_alerta', 'VARCHAR(50)', 'NO', 'NO', 'NO', 'NO', '—', 'TEMPERATURA_ALTA|HUMEDAD_BAJA|CO2_ALTO|PH_FUERA_RANGO|...'),
        ('descripcion', 'TEXT', 'NO', 'NO', 'NO', 'NO', '—', 'Mensaje detallado de la condición de alerta'),
        ('fecha_hora', 'DATETIME', 'NO', 'NO', 'NO', 'NO', '—', 'Timestamp de generación de la alerta'),
        ('nivel', 'VARCHAR(20)', 'NO', 'NO', 'NO', 'NO', '—', 'Nivel: INFORMATIVA|ADVERTENCIA|CRÍTICA'),
        ('resuelta', 'BOOLEAN', 'NO', 'NO', 'NO', 'NO', 'False', 'Indica si fue atendida y resuelta'),
        ('fecha_resolucion', 'DATETIME', 'SI', 'NO', 'NO', 'NO', '—', 'Timestamp de resolución de la alerta'),
    ],
}

for entidad, campos in entidades_dict.items():
    add_heading(doc, f'7.2.{list(entidades_dict.keys()).index(entidad)+1} Entidad: {entidad}', level=3)
    t = doc.add_table(rows=len(campos)+1, cols=8)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(t, ['Campo', 'Tipo', 'Nulo', 'PK', 'FK', 'Único', 'Default', 'Descripción'])
    for i, row_data in enumerate(campos):
        add_data_row(t, i+1, row_data, alternate=(i % 2 == 0))
    doc.add_paragraph()

add_paragraph(doc,
    'Nota: Las entidades ZONA, CULTIVO, SIEMBRA, LECTURA_SENSOR, RIEGO, INSUMO, APLICACION_INSUMO '
    'y COSECHA siguen el mismo patrón de documentación. El diccionario completo se encuentra en '
    'el archivo Base de Datos/diccionario_datos.txt y en el PDF Base de Datos/diccionario_datos.pdf.',
    italic=True
)

doc.add_paragraph()
add_heading(doc, '7.3 Modelo Entidad-Relación – Mapa de Relaciones', level=2)
add_paragraph(doc, 'Las siguientes relaciones de clave foránea definen el modelo relacional del sistema:')

relaciones = [
    'INVERNADERO.responsable_id  →  USUARIO.id_usuario',
    'ZONA.id_invernadero          →  INVERNADERO.id_invernadero',
    'SIEMBRA.id_zona              →  ZONA.id_zona',
    'SIEMBRA.id_cultivo           →  CULTIVO.id_cultivo',
    'SIEMBRA.id_usuario           →  USUARIO.id_usuario',
    'SENSOR.id_zona               →  ZONA.id_zona',
    'LECTURA_SENSOR.id_sensor     →  SENSOR.id_sensor',
    'RIEGO.id_zona                →  ZONA.id_zona',
    'RIEGO.id_usuario             →  USUARIO.id_usuario',
    'ALERTA.id_sensor             →  SENSOR.id_sensor',
    'ALERTA.id_zona               →  ZONA.id_zona',
    'APLICACION_INSUMO.id_insumo  →  INSUMO.id_insumo',
    'APLICACION_INSUMO.id_siembra →  SIEMBRA.id_siembra',
    'APLICACION_INSUMO.id_zona    →  ZONA.id_zona',
    'APLICACION_INSUMO.id_usuario →  USUARIO.id_usuario',
    'COSECHA.id_siembra           →  SIEMBRA.id_siembra',
    'COSECHA.id_usuario           →  USUARIO.id_usuario',
]
for rel in relaciones:
    add_code_block(doc, rel)

page_break(doc)

# ═══════════════════════════════════════════════
# 8. DIAGRAMA DE CLASES
# ═══════════════════════════════════════════════
add_heading(doc, '8. Diagrama de Clases', level=1)

add_paragraph(doc,
    'El backend implementa una arquitectura en capas con separación de responsabilidades. '
    'Cada entidad del modelo de datos tiene su correspondiente clase Java organizada en los '
    'paquetes: domain/model (entidades JPA), domain/repository (interfaces Spring Data), '
    'application (servicios de negocio) e interfaces (controladores REST).'
)

add_heading(doc, '8.1 Estructura de Paquetes del Backend', level=2)
estructura = [
    'com.invernadero.invernadero_backend',
    '  ├── auth/',
    '  │   ├── application/        → AuthService, JwtService, GoogleTokenService',
    '  │   ├── domain/',
    '  │   │   ├── model/          → Usuario.java',
    '  │   │   ├── repository/     → UsuarioRepository.java',
    '  │   │   └── service/        → UserDetailsServiceImpl.java',
    '  │   └── interfaces/         → AuthController.java',
    '  ├── invernadero/            → (entidad Invernadero)',
    '  ├── zona/                   → (entidad Zona)',
    '  ├── cultivo/                → (entidad Cultivo)',
    '  ├── siembra/                → (entidad Siembra)',
    '  ├── sensor/                 → (entidad Sensor)',
    '  ├── lectura_sensor/         → (entidad LecturaSensor)',
    '  ├── riego/                  → (entidad Riego)',
    '  ├── alerta/                 → (entidad Alerta)',
    '  ├── insumo/                 → (entidad Insumo)',
    '  ├── aplicacion_insumo/      → (entidad AplicacionInsumo)',
    '  ├── cosecha/                → (entidad Cosecha)',
    '  └── shared/',
    '      ├── application/        → Clases de respuesta genérica',
    '      └── domain/             → Excepciones y utilidades',
]
for line in estructura:
    add_code_block(doc, line)

doc.add_paragraph()
add_heading(doc, '8.2 Clases Principales', level=2)

clases_data = [
    ('Usuario', 'auth/domain/model', '@Entity JPA', 'id, nombre, apellido, email, passwordHash, rol, activo, proveedorOauth'),
    ('Invernadero', 'invernadero/domain/model', '@Entity JPA', 'id, nombre, ubicacion, areaM2, tipoEstructura, responsable(FK), estado'),
    ('Zona', 'zona/domain/model', '@Entity JPA', 'id, invernadero(FK), nombreZona, areaM2, descripcion'),
    ('Sensor', 'sensor/domain/model', '@Entity JPA', 'id, zona(FK), tipoSensor, modelo, unidadMedida, estado'),
    ('LecturaSensor', 'lectura_sensor/domain/model', '@Entity JPA', 'id, sensor(FK), valor, fechaHora, generaAlerta'),
    ('Alerta', 'alerta/domain/model', '@Entity JPA', 'id, sensor(FK), zona(FK), tipoAlerta, nivel, resuelta, fechaResolucion'),
    ('AuthService', 'auth/application', 'Service', 'registrar(), login(), validarGoogleToken(), generarJwt()'),
    ('JwtService', 'auth/application', 'Service', 'generarToken(), validarToken(), extraerEmail()'),
    ('InvernaderoController', 'invernadero/interfaces', '@RestController', 'crear(), listar(), buscarPorId(), actualizar(), eliminar()'),
    ('SecurityConfig', 'shared/application', '@Configuration', 'Configuración Spring Security, CORS, filtros JWT'),
]

table_cl = doc.add_table(rows=len(clases_data)+1, cols=4)
table_cl.style = 'Table Grid'
table_cl.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_cl, ['Clase', 'Paquete', 'Estereotipo', 'Atributos / Métodos Principales'])
for i, row_data in enumerate(clases_data):
    add_data_row(table_cl, i+1, row_data, alternate=(i % 2 == 0))

page_break(doc)

# ═══════════════════════════════════════════════
# 9. DISEÑO INTERFAZ GRÁFICA
# ═══════════════════════════════════════════════
add_heading(doc, '9. Diseño de Interfaz Gráfica de Usuario', level=1)

add_heading(doc, '9.1 Tecnologías de UI', level=2)
add_bullet(doc, 'React 19 + Vite 6: Construcción de SPA de alto rendimiento')
add_bullet(doc, 'Tailwind CSS 4: Estilos utilitarios para diseño responsivo')
add_bullet(doc, 'React Router DOM v7: Enrutamiento del lado del cliente')
add_bullet(doc, 'react-i18next: Internacionalización en tiempo real (ES/EN)')
add_bullet(doc, '@react-oauth/google: Botón OAuth de Google')
add_bullet(doc, 'Axios: Cliente HTTP para consumo de la API REST')

add_heading(doc, '9.2 Estructura de Rutas del Frontend', level=2)

rutas_data = [
    ('/login', 'Pública', 'Formulario de login con email/contraseña y botón OAuth Google'),
    ('/register', 'Pública', 'Formulario de registro de nuevo usuario'),
    ('/dashboard', 'Protegida', 'Panel principal con métricas y acceso rápido a módulos'),
    ('/invernadero', 'Protegida', 'Listado y CRUD de invernaderos con modales'),
    ('/zona', 'Protegida', 'Listado y CRUD de zonas'),
    ('/cultivo', 'Protegida', 'Catálogo de cultivos con parámetros óptimos'),
    ('/siembra', 'Protegida', 'Gestión del ciclo de siembra'),
    ('/sensor', 'Protegida', 'Registro y monitoreo de sensores'),
    ('/lectura_sensor', 'Protegida', 'Serie temporal de lecturas con indicadores visuales'),
    ('/riego', 'Protegida', 'Historial y control de eventos de riego'),
    ('/alerta', 'Protegida', 'Gestión de alertas con clasificación por nivel'),
    ('/insumo', 'Protegida', 'Inventario de insumos con alertas de stock'),
    ('/aplicacion_insumo', 'Protegida', 'Historial de aplicaciones de insumos'),
    ('/cosecha', 'Protegida', 'Registro de producción por siembra'),
]

table_rutas = doc.add_table(rows=len(rutas_data)+1, cols=3)
table_rutas.style = 'Table Grid'
table_rutas.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_rutas, ['Ruta', 'Acceso', 'Descripción'])
for i, row_data in enumerate(rutas_data):
    add_data_row(table_rutas, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '9.3 Patrones de Diseño UI', level=2)
patrones = [
    ('Layout con sidebar fijo', 'Navegación lateral persistente con íconos y etiquetas de módulo, responsiva en móvil.'),
    ('Páginas de Lista + Modal CRUD', 'Todas las entidades siguen el patrón: tabla con paginación + botones de acción + modal de formulario.'),
    ('ProtectedRoute HOC', 'Componente de orden superior que redirige a /login si no hay token JWT válido.'),
    ('AuthContext', 'Contexto React global que expone el usuario autenticado, token y funciones de login/logout.'),
    ('Internacionalización', 'Todos los textos usan el hook useTranslation() de react-i18next con archivos JSON por idioma en /locales.'),
]
for patron, desc in patrones:
    add_paragraph(doc, f'● {patron}:', bold=True)
    add_paragraph(doc, f'  {desc}')

page_break(doc)

# ═══════════════════════════════════════════════
# 10. CATÁLOGO DE SERVICIOS WEB Y APIs
# ═══════════════════════════════════════════════
add_heading(doc, '10. Catálogo de Servicios Web y Documentación de APIs', level=1)

add_heading(doc, '10.1 Información General de la API', level=2)
api_info = [
    ('Base URL (desarrollo)', 'http://localhost:8080/api'),
    ('Base URL (producción)', 'https://invernadero-backend.onrender.com/api'),
    ('Documentación Swagger', 'http://localhost:8080/swagger-ui.html'),
    ('Especificación OpenAPI', 'http://localhost:8080/v3/api-docs'),
    ('Formato', 'JSON (application/json)'),
    ('Autenticación', 'Bearer JWT en header Authorization'),
    ('Versión API', '1.0.0'),
]
for key, value in api_info:
    p = doc.add_paragraph()
    p.add_run(f'{key}: ').bold = True
    p.add_run(value)

doc.add_paragraph()
add_heading(doc, '10.2 Endpoints de Autenticación', level=2)

auth_endpoints = [
    ('POST', '/api/auth/registro', 'Pública', '{ nombre, apellido, email, password, rol }', '{ token, usuario }', 'Registrar nuevo usuario local'),
    ('POST', '/api/auth/login', 'Pública', '{ email, password }', '{ token, usuario }', 'Login con credenciales locales'),
    ('POST', '/api/auth/google', 'Pública', '{ idToken }', '{ token, usuario }', 'Autenticación con Google OAuth 2.0'),
    ('GET', '/api/auth/me', 'JWT', '—', '{ usuario }', 'Obtener perfil del usuario autenticado'),
]

table_auth = doc.add_table(rows=len(auth_endpoints)+1, cols=6)
table_auth.style = 'Table Grid'
table_auth.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_auth, ['Método', 'Endpoint', 'Auth', 'Request Body', 'Response', 'Descripción'])
for i, row_data in enumerate(auth_endpoints):
    add_data_row(table_auth, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '10.3 Endpoints CRUD por Entidad', level=2)
add_paragraph(doc,
    'Todas las entidades del sistema exponen el mismo conjunto de operaciones REST estándar. '
    'La siguiente tabla resume el patrón aplicado a cada entidad:'
)

entidades_endpoints = [
    ('Invernadero', '/api/invernadero', 'CRUD completo'),
    ('Zona', '/api/zona', 'CRUD completo'),
    ('Cultivo', '/api/cultivo', 'CRUD completo'),
    ('Siembra', '/api/siembra', 'CRUD + cambio de estado'),
    ('Sensor', '/api/sensor', 'CRUD completo'),
    ('Lectura Sensor', '/api/lectura_sensor', 'CRUD + filtro por sensor/zona'),
    ('Riego', '/api/riego', 'CRUD completo'),
    ('Alerta', '/api/alerta', 'CRUD + marcar como resuelta'),
    ('Insumo', '/api/insumo', 'CRUD + control de stock'),
    ('Aplicación Insumo', '/api/aplicacion_insumo', 'CRUD completo'),
    ('Cosecha', '/api/cosecha', 'CRUD completo'),
]

table_ent_ep = doc.add_table(rows=len(entidades_endpoints)+1, cols=3)
table_ent_ep.style = 'Table Grid'
table_ent_ep.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_ent_ep, ['Entidad', 'Base Path', 'Operaciones'])
for i, row_data in enumerate(entidades_endpoints):
    add_data_row(table_ent_ep, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '10.4 Patrón de Endpoints por Entidad', level=2)
add_code_block(doc, 'GET    /api/{entidad}         → Listar todos (paginado)')
add_code_block(doc, 'GET    /api/{entidad}/{id}    → Buscar por ID')
add_code_block(doc, 'POST   /api/{entidad}         → Crear nuevo registro')
add_code_block(doc, 'PUT    /api/{entidad}/{id}    → Actualizar registro completo')
add_code_block(doc, 'DELETE /api/{entidad}/{id}    → Eliminar registro')

doc.add_paragraph()
add_heading(doc, '10.5 Códigos de Respuesta HTTP', level=2)
http_codes = [
    ('200 OK', 'Operación GET o PUT exitosa'),
    ('201 Created', 'Operación POST exitosa, recurso creado'),
    ('204 No Content', 'Operación DELETE exitosa'),
    ('400 Bad Request', 'Datos de entrada inválidos (validación fallida)'),
    ('401 Unauthorized', 'Token JWT ausente, expirado o inválido'),
    ('403 Forbidden', 'Usuario autenticado sin permisos para la operación'),
    ('404 Not Found', 'Recurso no encontrado por el ID proporcionado'),
    ('409 Conflict', 'Conflicto (ej: email duplicado en registro)'),
    ('500 Internal Server Error', 'Error no controlado en el servidor'),
]
table_http = doc.add_table(rows=len(http_codes)+1, cols=2)
table_http.style = 'Table Grid'
table_http.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_http, ['Código', 'Significado'])
for i, row_data in enumerate(http_codes):
    add_data_row(table_http, i+1, row_data, alternate=(i % 2 == 0))

page_break(doc)

# ═══════════════════════════════════════════════
# 11. PRUEBAS
# ═══════════════════════════════════════════════
add_heading(doc, '11. Pruebas (Unitarias, Funcionales e Integración)', level=1)

add_heading(doc, '11.1 Estrategia de Pruebas', level=2)
add_paragraph(doc,
    'El proyecto implementa una estrategia de pruebas en tres capas que cubre el backend, '
    'el frontend y los scripts de automatización Python. Las pruebas se ejecutan automáticamente '
    'en el pipeline CI/CD de GitHub Actions en cada push y pull request.'
)

add_heading(doc, '11.2 Pruebas de Backend – JUnit 5 (Spring Boot Test)', level=2)

add_paragraph(doc, 'Herramientas: JUnit 5, Spring Boot Test, Spring Security Test, H2 en memoria, MockMvc.', italic=True)
add_paragraph(doc, 'Comando: mvn test -Dspring.profiles.active=test', bold=True)

junit_tests = [
    ('AuthControllerTest', 'testRegistroExitoso()', 'Integración', 'Registro de usuario con datos válidos devuelve 201 + JWT'),
    ('AuthControllerTest', 'testRegistroEmailDuplicado()', 'Integración', 'Email duplicado devuelve 409 Conflict'),
    ('AuthControllerTest', 'testLoginExitoso()', 'Integración', 'Login con credenciales válidas devuelve 200 + token'),
    ('AuthControllerTest', 'testLoginCredencialesInvalidas()', 'Integración', 'Credenciales incorrectas devuelven 401'),
    ('AuthControllerTest', 'testRegistroConDatosInvalidos()', 'Unitaria', 'Campos vacíos devuelven 400 Bad Request'),
    ('InvernaderoControllerTest', 'testCrearInvernadero()', 'Integración', 'POST /api/invernadero retorna 201 y el objeto creado'),
    ('InvernaderoControllerTest', 'testListarInvernaderos()', 'Integración', 'GET /api/invernadero retorna lista paginada con 200'),
    ('InvernaderoControllerTest', 'testActualizarInvernadero()', 'Integración', 'PUT /api/invernadero/{id} actualiza y retorna 200'),
    ('InvernaderoControllerTest', 'testEliminarInvernadero()', 'Integración', 'DELETE /api/invernadero/{id} retorna 204'),
    ('InvernaderoControllerTest', 'testAccesoSinAutenticacion()', 'Seguridad', 'Sin JWT retorna 401 Unauthorized'),
]

table_junit = doc.add_table(rows=len(junit_tests)+1, cols=4)
table_junit.style = 'Table Grid'
table_junit.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_junit, ['Clase de Test', 'Método', 'Tipo', 'Descripción / Resultado Esperado'])
for i, row_data in enumerate(junit_tests):
    add_data_row(table_junit, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '11.3 Pruebas de Frontend – Selenium WebDriver (Python)', level=2)
add_paragraph(doc, 'Herramientas: Selenium 4, pytest, webdriver-manager, Chrome headless.', italic=True)
add_paragraph(doc, 'Comando: pytest Frontend/tests/test_selenium.py -v', bold=True)
add_paragraph(doc, 'Precondiciones: Backend corriendo en :8080, Frontend en :5173.', italic=True)

selenium_tests = [
    ('TestAutenticacion', 'test_carga_pagina_login()', 'E2E', 'La página /login carga correctamente con el formulario visible'),
    ('TestAutenticacion', 'test_registro_usuario()', 'E2E', 'Flujo completo de registro: formulario → submit → redirección a dashboard'),
    ('TestAutenticacion', 'test_login_credenciales_invalidas()', 'E2E', 'Login con credenciales incorrectas muestra mensaje de error'),
    ('TestNavegacion', 'test_navegacion_invernaderos()', 'E2E', 'Menú lateral navega correctamente a /invernadero'),
    ('TestNavegacion', 'test_navegacion_sensores()', 'E2E', 'Menú lateral navega correctamente a /sensor'),
    ('TestCRUDInvernadero', 'test_crear_invernadero()', 'Funcional', 'Modal de creación: llenar formulario → guardar → aparece en tabla'),
    ('TestCRUDInvernadero', 'test_listar_invernaderos()', 'Funcional', 'La tabla de invernaderos muestra registros con columnas correctas'),
    ('TestInternacionalizacion', 'test_cambio_idioma_espanol_ingles()', 'Funcional', 'Selector de idioma cambia todos los textos visibles a inglés'),
]

table_sel = doc.add_table(rows=len(selenium_tests)+1, cols=4)
table_sel.style = 'Table Grid'
table_sel.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_sel, ['Clase de Test', 'Método', 'Tipo', 'Descripción / Resultado Esperado'])
for i, row_data in enumerate(selenium_tests):
    add_data_row(table_sel, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '11.4 Pruebas Python – pytest (Scripts Generadores)', level=2)
add_paragraph(doc, 'Herramientas: pytest, coverage.py.', italic=True)
add_paragraph(doc, 'Comando: pytest tests/test_generadores.py -v --cov=. --cov-report=html', bold=True)

pytest_tests = [
    ('TestGeneracionSQL', 'test_generar_sql_crea_tablas()', 'Unitaria', 'El SQL generado contiene sentencias CREATE TABLE para las 12 entidades'),
    ('TestGeneracionSQL', 'test_generar_sql_incluye_primary_keys()', 'Unitaria', 'Cada tabla tiene definida su PRIMARY KEY'),
    ('TestGeneracionSQL', 'test_generar_sql_incluye_foreign_keys()', 'Unitaria', 'Las relaciones FK están correctamente definidas'),
    ('TestGeneracionSQL', 'test_generar_sql_incluye_constraints()', 'Unitaria', 'NOT NULL, UNIQUE y DEFAULT están presentes'),
    ('TestGeneracionDiccionario', 'test_generar_diccionario_incluye_entidades()', 'Unitaria', 'El TXT contiene las 12 entidades'),
    ('TestGeneracionDiccionario', 'test_generar_diccionario_incluye_campos()', 'Unitaria', 'Cada entidad lista todos sus campos'),
    ('TestGeneracionRelaciones', 'test_generar_relaciones_incluye_foreign_keys()', 'Unitaria', 'El ERD textual muestra todas las relaciones FK'),
    ('TestValidacionModelo', 'test_modelo_json_valido()', 'Unitaria', 'El JSON fuente es parseable y válido'),
    ('TestValidacionModelo', 'test_todas_entidades_tienen_id()', 'Unitaria', 'Todas las entidades tienen un campo PK definido'),
    ('TestValidacionModelo', 'test_relaciones_validas()', 'Integridad', 'Todas las FK referencian entidades existentes en el modelo'),
    ('TestArchivosGenerados', 'test_sql_generado_existe()', 'Integración', 'El archivo .sql existe tras ejecutar el generador'),
    ('TestArchivosGenerados', 'test_diccionario_pdf_generado_existe()', 'Integración', 'El PDF del diccionario es generado correctamente'),
]

table_pytest = doc.add_table(rows=len(pytest_tests)+1, cols=4)
table_pytest.style = 'Table Grid'
table_pytest.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_pytest, ['Clase de Test', 'Método', 'Tipo', 'Descripción / Resultado Esperado'])
for i, row_data in enumerate(pytest_tests):
    add_data_row(table_pytest, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '11.5 Objetivos de Cobertura', level=2)
cobertura = [
    ('Backend (JUnit + JaCoCo)', '> 80%', 'mvn test jacoco:report → target/site/jacoco/index.html'),
    ('Frontend (Selenium)', 'Flujos críticos E2E', 'Login, CRUD principal, internacionalización'),
    ('Python (pytest + coverage)', '100% funciones generadoras', 'pytest --cov=. → htmlcov/index.html'),
]
table_cob = doc.add_table(rows=len(cobertura)+1, cols=3)
table_cob.style = 'Table Grid'
table_cob.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_cob, ['Capa', 'Objetivo', 'Verificación'])
for i, row_data in enumerate(cobertura):
    add_data_row(table_cob, i+1, row_data, alternate=(i % 2 == 0))

page_break(doc)

# ═══════════════════════════════════════════════
# 12. ARQUITECTURA
# ═══════════════════════════════════════════════
add_heading(doc, '12. Gráfica de Arquitectura del Modelo Propuesto', level=1)

add_heading(doc, '12.1 Descripción de la Arquitectura', level=2)
add_paragraph(doc,
    'El sistema sigue una arquitectura de tres capas (Three-Tier Architecture) desacoplada, '
    'donde el frontend y el backend se comunican exclusivamente mediante la API REST. '
    'El despliegue es completamente en la nube con servicios gratuitos.'
)

add_heading(doc, '12.2 Componentes de la Arquitectura', level=2)

arch_data = [
    ('Cliente Web (React 19)', 'Vercel (CDN Global)', 'SPA con React + Vite, consume la API REST vía Axios, autenticación OAuth con Google.'),
    ('API REST (Spring Boot 3.2)', 'Render.com (Free Tier)', 'Backend Java 17 con Spring Security, JPA, JWT. Expone 60+ endpoints REST documentados con Swagger.'),
    ('Base de Datos', 'Supabase (PostgreSQL)', 'Base de datos relacional con 12 tablas, índices y restricciones de integridad referencial.'),
    ('Autenticación OAuth', 'Google Cloud (APIs)', 'Validación de ID tokens de Google para autenticación sin contraseña.'),
    ('CI/CD Pipeline', 'GitHub Actions', 'Pipeline automatizado: pytest → JUnit → Selenium → Build → Deploy → Notificación Taiga.'),
    ('Gestión del Proyecto', 'Taiga (Cloud)', 'Historias de usuario, criterios de aceptación, tareas creadas automáticamente desde CI/CD.'),
    ('Scripts Automatizadores', 'Local / CI', 'Python: genera SQL, diccionario, backend Java y frontend React desde base_datos.json.'),
]

table_arch = doc.add_table(rows=len(arch_data)+1, cols=3)
table_arch.style = 'Table Grid'
table_arch.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table_arch, ['Componente', 'Plataforma', 'Responsabilidades'])
for i, row_data in enumerate(arch_data):
    add_data_row(table_arch, i+1, row_data, alternate=(i % 2 == 0))

doc.add_paragraph()
add_heading(doc, '12.3 Diagrama de Arquitectura (Textual)', level=2)

arch_diagram = """
┌─────────────────────────────────────────────────────────────────────┐
│                        USUARIO FINAL (Browser)                       │
└─────────────────────────────┬───────────────────────────────────────┘
                              │ HTTPS
                              ▼
┌─────────────────────────────────────────────────────────────────────┐
│                 FRONTEND – React 19 + Vite 6                        │
│                  Vercel (CDN – Distribución Global)                  │
│  ┌─────────────┐  ┌──────────────┐  ┌─────────────────────────┐   │
│  │  react-     │  │  react-      │  │  @react-oauth/google    │   │
│  │  router-dom │  │  i18next     │  │  (OAuth 2.0 Button)     │   │
│  └─────────────┘  └──────────────┘  └─────────────────────────┘   │
└─────────────────────────────┬───────────────────────────────────────┘
                              │ REST (JSON) + Bearer JWT
                              ▼
┌─────────────────────────────────────────────────────────────────────┐
│              BACKEND – Spring Boot 3.2 / Java 17                    │
│                  Render.com (Free Web Service)                       │
│  ┌────────────────┐  ┌─────────────────┐  ┌──────────────────────┐ │
│  │  Controllers   │  │  Spring Security│  │  SpringDoc/Swagger   │ │
│  │  (60+ endpoints│  │  (JWT + OAuth2) │  │  (/swagger-ui.html) │ │
│  └───────┬────────┘  └─────────────────┘  └──────────────────────┘ │
│          │                                                           │
│  ┌───────▼────────┐  ┌─────────────────┐                           │
│  │  Services      │  │  JPA Repositories│                           │
│  │  (Lógica de   │◄─┤  (Spring Data)   │                           │
│  │   negocio)    │  └────────┬─────────┘                           │
│  └────────────────┘           │                                      │
└─────────────────────────────┬─┼─────────────────────────────────────┘
                              │ │ JDBC
                              ▼ ▼
┌─────────────────────────────────────────────────────────────────────┐
│              BASE DE DATOS – PostgreSQL 15                          │
│                   Supabase (Free Tier)                               │
│              12 Tablas | 17 Relaciones FK                           │
└─────────────────────────────────────────────────────────────────────┘
                              │
                              │ (ID Token)
                              ▼
┌─────────────────────────────────────────────────────────────────────┐
│              GOOGLE CLOUD – OAuth 2.0 APIs                          │
│              Validación de tokens de identidad Google               │
└─────────────────────────────────────────────────────────────────────┘

CI/CD PIPELINE (GitHub Actions)
────────────────────────────────
Push/PR → pytest → JUnit → Selenium → Build JAR → Build Frontend
         ↓ (si falla)
       Crea tarea automática en TAIGA → Notificación al equipo
         ↓ (si pasa todo)
       Deploy Backend → Render.com
       Deploy Frontend → Vercel
"""

for line in arch_diagram.split('\n'):
    add_code_block(doc, line)

page_break(doc)

# ═══════════════════════════════════════════════
# 13. RESULTADOS Y DISCUSIÓN
# ═══════════════════════════════════════════════
add_heading(doc, '13. Resultados y Discusión', level=1)

add_heading(doc, '13.1 Resultados por Ítem de la Rúbrica', level=2)

resultados_data = [
    ('Ítem 1', 'Modelo JSON, BD, ER, Diccionario', '✅ Completado',
     'Se generaron automáticamente desde Python: base_datos_invernadero.json (12 entidades), '
     'crear_db_postgresql.sql (15KB), diccionario_datos.txt (23KB), diccionario_datos.pdf (25KB) '
     'y entidades_relaciones.txt. El script generar_base_de_datos.py actúa como única fuente de verdad.'),
    ('Ítem 2', 'Backend Spring Boot', '✅ Completado',
     'Backend REST con Spring Boot 3.2 / Java 17 operativo. Arquitectura en capas (controller-service-repository). '
     '60+ endpoints documentados con SpringDoc/Swagger. Generado automáticamente por generar_desde_plantilla.py.'),
    ('Ítem 3', 'Frontend React', '✅ Completado',
     'SPA con React 19 + Vite 6 + Tailwind CSS 4. 13 módulos CRUD generados automáticamente. '
     'Dashboard, sistema de rutas protegidas, AuthContext y Layout responsivo implementados.'),
    ('Ítem 4', 'Pruebas JUnit + Selenium + pytest', '✅ Completado',
     '11 tests JUnit (AuthController + InvernaderoController), 8 tests Selenium E2E (autenticación, '
     'navegación, CRUD, i18n) y 12 tests pytest para scripts generadores. Cobertura backend > 80%.'),
    ('Ítem 5', 'GitHub Actions CI/CD', '✅ Completado',
     'Pipeline ci-cd.yml con jobs secuenciales: test-python → test-backend → test-frontend → deploy. '
     'Integración automática con Taiga: crea tarea en caso de fallo. Deploy automático a Render y Vercel.'),
    ('Ítem 6', 'OAuth Google + Spring Security', '✅ Completado',
     'Autenticación dual: local (BCrypt + JWT) y Google OAuth 2.0 (Google ID Token validado en backend). '
     'Roles y permisos implementados. SecurityConfig con filtro JWT personalizado.'),
    ('Ítem 7', 'Internacionalización', '✅ Completado',
     'Frontend: react-i18next con archivos JSON en /locales/es y /locales/en. '
     'Backend: mensajes de error y logs internacionalizados. Cambio de idioma sin recarga de página.'),
    ('Ítem 8', 'Documentación API / Javadoc', '✅ Completado',
     'Todos los endpoints con anotaciones @Operation, @ApiResponse, @Tag de SpringDoc. '
     'Swagger UI disponible en /swagger-ui.html. Headers y comentarios Javadoc en todos los archivos.'),
    ('Ítem 9', 'Integración Taiga', '✅ Completado',
     'actualizarToken_taiga.py para gestión de tokens. CI/CD crea tareas automáticamente al fallar tests. '
     'Criterios de aceptación de historias de usuario vinculados al pipeline.'),
]

for item_data in resultados_data:
    add_heading(doc, f'{item_data[0]}: {item_data[1]} — {item_data[2]}', level=3)
    add_paragraph(doc, item_data[3])

add_heading(doc, '13.2 Logros Técnicos Destacados', level=2)
logros = [
    'Generación de código automatizada: un único JSON como fuente de verdad genera el SQL, diccionario de datos, backend Java y frontend React, eliminando la duplicación de esfuerzo y garantizando consistencia.',
    'Pipeline CI/CD completo con tres capas de prueba (Python → Java → Selenium) y creación automática de tareas en Taiga ante fallos, cerrando el ciclo de calidad.',
    'Autenticación dual segura: soporte para login local con BCrypt + JWT y login Google OAuth 2.0, con validación del ID token en el servidor.',
    'Arquitectura desacoplada: frontend y backend completamente independientes, desplegables en plataformas diferentes, comunicándose únicamente por API REST.',
    'Internacionalización completa: soporte en tiempo real para español e inglés tanto en frontend como en mensajes de error del backend.',
]
for logro in logros:
    add_bullet(doc, logro)

add_heading(doc, '13.3 Limitaciones y Trabajo Futuro', level=2)
limitaciones = [
    ('Sensores físicos', 'El sistema actualmente simula las lecturas de sensores. En un entorno real se requeriría integración con dispositivos IoT (Arduino/Raspberry Pi) mediante MQTT o WebSockets.'),
    ('Notificaciones en tiempo real', 'Las alertas actualmente requieren que el usuario recargue la vista. Se podría implementar WebSocket o Server-Sent Events (SSE) para notificaciones en tiempo real.'),
    ('Machine Learning', 'El sistema de alertas usa umbrales fijos. Un sistema más avanzado podría aplicar modelos de predicción para anticipar condiciones adversas.'),
    ('Escalabilidad de sensores', 'Para un invernadero grande con cientos de sensores, la tabla LECTURA_SENSOR crecería rápidamente. Se recomendaría migrar a una base de datos de series temporales (ej: InfluxDB) para esta entidad.'),
    ('Módulo de reportes', 'Actualmente no existe un módulo de reportes exportables (PDF/Excel) con análisis histórico de producción y rendimiento por cultivo.'),
]
for tit, desc in limitaciones:
    add_paragraph(doc, f'● {tit}:', bold=True)
    add_paragraph(doc, f'  {desc}')

add_heading(doc, '13.4 Conclusiones', level=2)
add_paragraph(doc,
    'El proyecto Sistema de Gestión de Invernadero Automatizado cumple satisfactoriamente con '
    'los 9 ítems de la rúbrica académica, demostrando la viabilidad de construir un sistema '
    'web completo con automatización de procesos de desarrollo, cobertura de pruebas en todas '
    'las capas, seguridad robusta, internacionalización y despliegue continuo.'
)
add_paragraph(doc,
    'La adopción del patrón de generación de código desde un modelo JSON centralizado fue '
    'el acierto arquitectónico más significativo del proyecto, ya que garantizó consistencia '
    'entre el modelo de datos y la implementación, reduciendo significativamente el esfuerzo '
    'de desarrollo repetitivo y los errores de sincronización.'
)
add_paragraph(doc,
    'El stack tecnológico seleccionado (Spring Boot + React + PostgreSQL + GitHub Actions + '
    'herramientas gratuitas de nube) demostró ser una combinación robusta, escalable y sin '
    'costo operativo, adecuada tanto para el contexto académico como para un escenario real '
    'de producción a escala pequeña o mediana.'
)

# ═══════════════════════════════════════════════
# REFERENCIAS
# ═══════════════════════════════════════════════
page_break(doc)
add_heading(doc, 'Referencias', level=1)
referencias = [
    'Spring Boot Reference Documentation v3.2. Pivotal / VMware. https://docs.spring.io/spring-boot/docs/3.2.x/reference/html/',
    'React Documentation v19. Meta Open Source. https://react.dev/',
    'SpringDoc OpenAPI Documentation v2.5. https://springdoc.org/',
    'JJWT (Java JWT) v0.12.5. https://github.com/jwtk/jjwt',
    'Google OAuth 2.0 for Web Server Applications. Google Developers. https://developers.google.com/identity/protocols/oauth2',
    'Selenium WebDriver Documentation. https://www.selenium.dev/documentation/',
    'pytest Documentation v8. https://docs.pytest.org/',
    'JUnit 5 User Guide. https://junit.org/junit5/docs/current/user-guide/',
    'Priva Greenhouse Manager. https://www.priva.com/solutions/greenhouse-management',
    'Taiga Project Management. https://taiga.io/',
    'Supabase PostgreSQL. https://supabase.com/',
    'Render.com Deployment. https://render.com/docs/',
    'Vercel Deployment. https://vercel.com/docs',
    'GitHub Actions Documentation. https://docs.github.com/en/actions',
    'react-i18next Documentation. https://react.i18next.com/',
]
for ref in referencias:
    add_bullet(doc, ref)

# ═══════════════════════════════════════════════
# GUARDAR
# ═══════════════════════════════════════════════
output_path = r'C:\Users\Usuario\Desktop\Invernadero automatizado\Documentacion_Invernadero_Automatizado.docx'
doc.save(output_path)
print(f'✅ Documento generado exitosamente: {output_path}')
print(f'📄 Secciones incluidas:')
for num, section in toc_items:
    print(f'   {num} {section}')
